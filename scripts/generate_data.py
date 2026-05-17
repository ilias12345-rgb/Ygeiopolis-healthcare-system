"""Generate deterministic CSV data for the Ygeiopolis hospital schema.

The generator combines cleaned official reference files, when available, with
synthetic transactional data designed to exercise the schema constraints,
triggers, views, and final reporting queries.
"""

from __future__ import annotations
import argparse
import json
import math
import random
import re
from collections import Counter, defaultdict
from dataclasses import dataclass
from datetime import date, datetime, time, timedelta
from pathlib import Path
import shutil
import subprocess
import tempfile

import numpy as np
import pandas as pd
try:
    from docx import Document
except Exception:
    Document = None

SEED = 42
random.seed(SEED)
np.random.seed(SEED)
DATASET_AS_OF_TS = datetime(2026, 5, 12, 12, 0, 0)
DATASET_START_DATE = date(2023, 5, 13)
DATASET_END_DATE = date(2026, 5, 12)
DEFAULT_SHIFT_DAYS = (DATASET_END_DATE - DATASET_START_DATE).days + 1
DEFAULT_SHIFT_SAMPLE_DAYS_PER_MONTH = 7

DEPARTMENT_SPECS = [
    ("Cardiology", "CARDIOLOGY"),
    ("Surgery", "GENERAL_SURGERY"),
    ("ICU", "INTENSIVE_CARE"),
    ("Emergency", "EMERGENCY_MEDICINE"),
    ("Neurology", "NEUROLOGY"),
    ("Orthopedics", "ORTHOPEDICS"),
    ("Pulmonology", "PULMONOLOGY"),
    ("Gastroenterology", "GASTROENTEROLOGY"),
    ("Oncology", "ONCOLOGY"),
    ("Pediatrics", "PEDIATRICS"),
    ("Ophthalmology", "OPHTHALMOLOGY"),
    ("ENT", "OTORHINOLARYNGOLOGY"),
    ("Nephrology", "NEPHROLOGY"),
    ("Psychiatry", "PSYCHIATRY"),
    ("Internal Medicine", "INTERNAL_MEDICINE"),
]

INSURANCE_PROVIDERS = ["EFKA", "PRIVATE_A", "PRIVATE_B", "UNINSURED", "ARMED_FORCES"]
NATIONALITIES = ["Greek", "Cypriot", "Albanian", "Bulgarian", "Romanian", "Italian", "German"]
PROFESSIONS = ["Teacher", "Engineer", "Driver", "Retired", "Student", "Farmer", "Clerk", "Lawyer", "Nurse", "Technician"]
ADMIN_ROLES = ["SECRETARY", "ACCOUNTANT", "REGISTRAR", "ARCHIVIST", "BILLING_OFFICER", "HR_ASSISTANT"]
OFFICE_WORKS = ["Front desk", "Billing office", "Medical records", "Scheduling", "HR office", "Admissions office"]
NURSE_RANKS = ["ASSISTANT_NURSE", "NURSE", "HEAD_NURSE"]
SHIFT_TIMES = {
    "MORNING": ("07:00:00", "15:00:00"),
    "AFTERNOON": ("15:00:00", "23:00:00"),
    "NIGHT": ("23:00:00", "07:00:00"),
}
BED_TYPES = ["ICU", "SINGLE", "MULTI_BED"]
BED_STATUSES = ["AVAILABLE", "AVAILABLE", "AVAILABLE", "OCCUPIED", "MAINTENANCE"]
PLACE_TYPES = ["OPERATING_ROOM"] * 8 + ["PROCEDURE_ROOM"] * 4

COMMON_SUBSTANCE_PAIRS = [
    ("PARACETAMOL", "ONDANSETRON"),
    ("AMOXICILLIN", "OMEPRAZOLE"),
    ("ENOXAPARIN", "CLOPIDOGREL"),
]

SYMPTOMS_BY_LEVEL = {
    1: ["Severe chest pain", "Loss of consciousness", "Major trauma", "Severe respiratory distress"],
    2: ["Stroke-like symptoms", "High fever with confusion", "Acute abdominal pain", "GI bleeding"],
    3: ["Fracture pain", "Dyspnoea", "Severe headache", "Moderate dehydration"],
    4: ["Back pain", "Mild asthma exacerbation", "Persistent cough", "Vomiting"],
    5: ["Medication refill issue", "Mild rash", "Chronic dizziness", "Non-urgent ear pain"],
}

FIRST_NAMES_M = ["Giorgos","Nikos","Dimitris","Kostas","Giannis","Panagiotis","Vasilis","Petros","Spyros","Andreas","Marios","Stavros","Theodoros","Michalis","Christos"]
FIRST_NAMES_F = ["Maria","Eleni","Katerina","Sofia","Georgia","Panagiota","Vasiliki","Anastasia","Dimitra","Ioanna","Despoina","Anna","Natalia","Chrysoula","Evangelia"]
LAST_NAMES = ["Papadopoulos","Nikolaou","Georgiou","Dimitriou","Vasiliou","Petrou","Ioannou","Kostopoulos","Pavlidis","Konstantinou","Athanasiou","Theodorou","Sotiriou","Lazarou","Alexiou"]
FATHER_NAMES = ["Ioannis","Nikolaos","Dimitrios","Konstantinos","Georgios","Panagiotis","Anastasios","Christos","Vasileios","Antonios"]

DIAG_KEYWORDS = ['ΔΙΑΓΝΩ','ΕΞΕΤΑΣ','ΑΚΤΙΝ','ΤΟΜΟΓ','ΥΠΕΡΗΧ','ΜΑΓΝΗΤ','ΗΛΕΚΤΡΟ','ΕΝΔΟΣΚΟΠ','ΑΓΓΕΙΟΓΡΑΦ','ΒΙΟΨ','ΣΠΙΡΟΜΕΤ','ΗΛΕΚΤΡΟΕΓΚΕΦ','ΚΑΘΕΤΗΡΙΑΣΜ','ΗΧΟΚΑΡΔ']
THER_KEYWORDS = ['ΘΕΡΑΠ','ΑΚΤΙΝΟΘΕΡ','ΧΗΜΕΙΟΘΕΡ','ΑΙΜΟΚΑΘΑΡ','ΔΙΑΛΥΣ','ΦΥΣΙΚΟΘΕΡ','ΜΕΤΑΓΓΙΣ','ΠΛΑΣΜΑΦΑΙΡ','ΕΜΦΥΤΕΥΣΗ ΣΥΣΚΕΥ']
SURG_KEYWORDS = ['ΕΠΕΜΒ','ΧΕΙΡΟΥΡ','ΑΦΑΙΡΕΣ','ΕΚΤΟΜ','ΣΥΡΡΑΦ','ΛΑΠΑΡ','ΑΡΘΡΟΣΚ','ΚΡΑΝΙΟ','ΑΚΡΩΤΗΡ','ΣΚΩΛΗΚΟΕΙΔΕΚ','ΧΟΛΟΚΥΣΤΕΚ']


def normalize_space(s):
    if s is None or (isinstance(s, float) and pd.isna(s)):
        return None
    s = str(s).replace("\xa0", " ").replace("\n", " ").replace("\r", " ")
    s = re.sub(r"\s+", " ", s).strip()
    return s or None


def parse_money_eur(s):
    if isinstance(s, (int, float, np.integer, np.floating)) and not pd.isna(s):
        return float(s)
    s = normalize_space(s)
    if not s:
        return None
    s = s.replace("€", "").replace("EUR", "").strip()
    s = s.replace(".", "").replace(",", ".")
    s = re.sub(r"[^0-9.]", "", s)
    try:
        return float(s) if s else None
    except ValueError:
        return None


def derive_proc_category(name: str) -> str:
    n = name.upper()
    if any(k in n for k in DIAG_KEYWORDS):
        return "DIAGNOSTIC"
    if any(k in n for k in THER_KEYWORDS):
        return "THERAPEUTIC"
    if any(k in n for k in SURG_KEYWORDS):
        return "SURGICAL"
    if "ΑΝΑΙΣΘΗΣ" in n or "ΤΟΜΗ" in n:
        return "SURGICAL"
    return "SURGICAL"


def derive_place_type(name: str, category: str) -> str:
    n = name.upper()
    if category == "SURGICAL":
        return "OPERATING_ROOM"
    if any(k in n for k in ["ΤΟΜΟΓ","ΜΑΓΝΗΤ","ΑΚΤΙΝ","ΥΠΕΡΗΧ","ΕΝΔΟΣΚΟΠ","ΑΓΓΕΙΟΓΡΑΦ"]):
        return "PROCEDURE_ROOM"
    return "PROCEDURE_ROOM"


def derive_standard_duration_min(name: str, category: str) -> int:
    seed = sum(ord(ch) for ch in name)
    if category == "SURGICAL":
        return 60 + (seed % 7) * 15
    if category == "THERAPEUTIC":
        return 30 + (seed % 5) * 15
    return 15 + (seed % 4) * 15


def derive_standard_cost(name: str, category: str) -> float:
    seed = sum(ord(ch) for ch in name)
    if category == "SURGICAL":
        return round(650.0 + (seed % 12) * 125.0, 2)
    if category == "THERAPEUTIC":
        return round(250.0 + (seed % 8) * 75.0, 2)
    return round(45.0 + (seed % 6) * 35.0, 2)


def derive_test_type(name: str) -> str:
    n = name.upper()
    if any(k in n for k in ["ΑΙΜΑ", "ΑΙΜΑΤ", "ΒΙΟΧΗΜ", "ΟΡΡ"]):
        return "HEMATOLOGY"
    if any(k in n for k in ["ΑΚΤΙΝ", "ΤΟΜΟΓ", "ΥΠΕΡΗΧ", "ΜΑΓΝΗΤ", "ΗΧΟΚΑΡΔ", "ΑΓΓΕΙΟΓΡΑΦ"]):
        return "IMAGING"
    if any(k in n for k in ["ΚΑΛΛΙΕΡΓ", "ΜΙΚΡΟΒ", "ΙΟΛ", "ΒΑΚΤΗΡ"]):
        return "MICROBIOLOGY"
    if any(k in n for k in ["ΒΙΟΨ", "ΙΣΤΟΛ"]):
        return "PATHOLOGY"
    return "DIAGNOSTIC_OTHER"


def pick(seq):
    return random.choice(list(seq))


def date_range(start: date, days: int):
    return [start + timedelta(days=i) for i in range(days)]


def mk_email(first, last, domain="yg-eupolis.gr", uniq=None):
    base = f"{first.lower()}.{last.lower()}".replace(" ", "").replace("'", "")
    if uniq is not None:
        base += str(uniq)
    return f"{base}@{domain}"


def phone(prefix="69"):
    return prefix + "".join(random.choice("0123456789") for _ in range(8))


def unique_amka(start_num: int) -> str:
    return f"{start_num:011d}"


class BedAllocator:
    """Track bed usage while synthetic hospitalizations are being created."""

    def __init__(self, beds_df: pd.DataFrame):
        self.by_dept = defaultdict(list)
        for row in beds_df.itertuples(index=False):
            self.by_dept[row.department_id].append(row.bed_id)
        self.busy = {bed_id: [] for bed_id in beds_df["bed_id"]}

    @staticmethod
    def overlaps(a_start, a_end, b_start, b_end):
        a_end = a_end or datetime.max
        b_end = b_end or datetime.max
        return a_start < b_end and b_start < a_end

    def allocate(self, department_id, start_ts, end_ts):
        for bed_id in self.by_dept[department_id]:
            if all(not self.overlaps(start_ts, end_ts, s, e) for s, e in self.busy[bed_id]):
                self.busy[bed_id].append((start_ts, end_ts))
                return bed_id
        return None


def resolve_source_file(source_dir: Path, candidates):
    search_roots = [source_dir, source_dir / "docs", source_dir / "converted"]
    normalized_candidates = [c for c in candidates if c]
    for root in search_roots:
        for cand in normalized_candidates:
            p = root / cand
            if p.exists():
                return p
    # recursive fallback by exact filename
    for root in search_roots:
        if root.exists():
            for cand in normalized_candidates:
                matches = list(root.rglob(Path(cand).name))
                if matches:
                    return matches[0]
    # relaxed fallback by stem prefix
    stems = [Path(c).stem for c in normalized_candidates]
    for root in search_roots:
        if root.exists():
            for file in root.rglob('*'):
                if file.is_file() and any(file.stem.startswith(stem) for stem in stems):
                    return file
    raise FileNotFoundError(f"Could not find any of: {normalized_candidates} under {source_dir}")


def read_word_table_or_text(path: Path):
    suffix = path.suffix.lower()
    if suffix == '.docx':
        if Document is None:
            raise FileNotFoundError(f'python-docx is unavailable for {path}')
        return Document(path), 'docx'
    if suffix == '.doc':
        soffice = shutil.which('soffice') or shutil.which('libreoffice')
        if soffice:
            with tempfile.TemporaryDirectory() as td:
                subprocess.run([soffice, '--headless', '--convert-to', 'docx', '--outdir', td, str(path)], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                converted = next(Path(td).glob('*.docx'))
                return Document(converted), 'docx'
        antiword = shutil.which('antiword')
        if antiword:
            txt = subprocess.run([antiword, str(path)], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True).stdout
            return txt, 'text'
        textutil = shutil.which('textutil')
        if textutil and Document is not None:
            with tempfile.TemporaryDirectory() as td:
                converted = Path(td) / f"{path.stem}.docx"
                subprocess.run([textutil, '-convert', 'docx', '-output', str(converted), str(path)], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                return Document(converted), 'docx'
    raise FileNotFoundError(f"Unsupported Word source file: {path}")


def looks_like_ken_code(value: str) -> bool:
    value = normalize_space(value)
    if not value or value.startswith('ΤΚΑ') or 'ΚΩΔΙΚ' in value:
        return False
    return bool(re.fullmatch(r"[A-ZΑ-Ω][A-ZΑ-Ω0-9-]{1,12}", value.upper()))


def parse_ken_from_lines(lines: list[str]) -> pd.DataFrame:
    rows = []
    cleaned = [normalize_space(line) for line in lines]
    cleaned = [line for line in cleaned if line]
    i = 0
    while i < len(cleaned):
        code = cleaned[i]
        if not looks_like_ken_code(code):
            i += 1
            continue

        desc_parts = []
        money = None
        days = None
        j = i + 1
        while j < len(cleaned):
            current = cleaned[j]
            current_money = parse_money_eur(current)
            if current_money is not None and re.search(r"\d", current):
                money = current_money
                j += 1
                break
            if looks_like_ken_code(current) and desc_parts:
                break
            desc_parts.append(current)
            j += 1

        if money is not None and j < len(cleaned):
            day_match = re.search(r"\d+", cleaned[j])
            if day_match:
                days = int(day_match.group(0))
                j += 1

        if desc_parts and money is not None and days is not None:
            rows.append((code.upper(), " ".join(desc_parts), round(money, 2), days))
            i = j
        else:
            i += 1

    return normalize_ken_dataframe(pd.DataFrame(rows, columns=['ken_code','ken_description','basic_cost','mean_duration_days']))


def parse_ken_from_text(txt: str):
    rows = []
    lines = [normalize_space(x) for x in txt.splitlines()]
    sequence_rows = parse_ken_from_lines(lines)
    if not sequence_rows.empty:
        return sequence_rows
    for line in lines:
        if not line:
            continue
        m = re.match(r'^([A-ZΑ-Ω][A-Z0-9\-]{1,10})\s+(.*)$', line)
        if m and not line.startswith('ΤΚΑ') and 'ΚΩΔΙΚ' not in line:
            code = m.group(1)
            rest = m.group(2)
            money = re.search(r'(\d{1,3}(?:\.\d{3})*(?:,\d+)?)\s*€', rest)
            mdn = re.search(r'(\d+)\s*$', rest)
            if money and mdn:
                desc = rest[:money.start()].strip(' -')
                rows.append((code, desc, round(parse_money_eur(money.group(0)),2), int(mdn.group(1))))
    return pd.DataFrame(rows, columns=['ken_code','ken_description','basic_cost','mean_duration_days']).drop_duplicates('ken_code').reset_index(drop=True)


def read_excel_flexible(path: Path, **kwargs):
    suffix = path.suffix.lower()
    engine = kwargs.pop('engine', None)
    if suffix == '.xls' and engine is None:
        engine = 'xlrd'
    return pd.read_excel(path, engine=engine, **kwargs)


def normalize_ken_dataframe(raw: pd.DataFrame) -> pd.DataFrame:
    """Normalize official or fallback KEN exports to the schema column order."""

    df = raw.copy()
    normalized_columns = {str(c).strip().lower(): c for c in df.columns}
    if "ken_code" in normalized_columns:
        rename = {
            normalized_columns["ken_code"]: "ken_code",
            normalized_columns.get("ken_description", "ken_description"): "ken_description",
            normalized_columns.get("basic_cost", "basic_cost"): "basic_cost",
            normalized_columns.get("mean_duration_days", "mean_duration_days"): "mean_duration_days",
            normalized_columns.get("extra_daily_cost", "extra_daily_cost"): "extra_daily_cost",
        }
        df = df.rename(columns={k: v for k, v in rename.items() if k in df.columns})
    else:
        cols = list(df.columns)
        rename = {
            cols[0]: "ken_code",
            cols[1]: "ken_description",
            cols[2]: "basic_cost",
            cols[3]: "mean_duration_days",
        }
        if len(cols) > 4:
            rename[cols[4]] = "extra_daily_cost"
        df = df.rename(columns=rename)

    required = ["ken_code", "ken_description", "basic_cost", "mean_duration_days"]
    df = df[[c for c in required + ["extra_daily_cost"] if c in df.columns]].copy()
    df["ken_code"] = df["ken_code"].map(normalize_space).str.upper()
    df["ken_description"] = df["ken_description"].map(normalize_space)
    df["basic_cost"] = pd.to_numeric(df["basic_cost"].map(parse_money_eur), errors="coerce")
    df["mean_duration_days"] = pd.to_numeric(df["mean_duration_days"], errors="coerce")
    df = df[
        df["ken_code"].notna()
        & df["ken_description"].notna()
        & df["basic_cost"].notna()
        & df["mean_duration_days"].notna()
    ].copy()
    df = df[~df["ken_code"].str.lower().eq("ken_code")]
    df = df[df["ken_code"].str.match(r"^[A-ZΑ-Ω][A-ZΑ-Ω0-9-]{1,12}$", na=False)]
    df["basic_cost"] = df["basic_cost"].round(2)
    df["mean_duration_days"] = df["mean_duration_days"].astype(int).clip(lower=1)
    if "extra_daily_cost" in df.columns:
        df["extra_daily_cost"] = pd.to_numeric(df["extra_daily_cost"].map(parse_money_eur), errors="coerce")
    else:
        df["extra_daily_cost"] = np.nan
    missing_extra = df["extra_daily_cost"].isna()
    df.loc[missing_extra, "extra_daily_cost"] = (
        df.loc[missing_extra, "basic_cost"] / df.loc[missing_extra, "mean_duration_days"]
    ).round(2)
    return df[["ken_code", "ken_description", "basic_cost", "mean_duration_days", "extra_daily_cost"]].drop_duplicates("ken_code").reset_index(drop=True)


def is_placeholder_ken(df: pd.DataFrame) -> bool:
    return not df.empty and df["ken_code"].astype(str).str.startswith("DKEN").all()


def parse_ken_from_docx(document) -> pd.DataFrame:
    rows = []
    for table in document.tables:
        for row in table.rows:
            cells = [normalize_space(cell.text) for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if len(cells) < 4:
                continue
            code = cells[0]
            if not re.match(r"^[A-ZΑ-Ω][A-ZΑ-Ω0-9-]{1,12}$", code or ""):
                continue
            numeric_cells = [cell for cell in cells if parse_money_eur(cell) is not None]
            day_cells = [cell for cell in cells if re.fullmatch(r"\d+", str(cell))]
            if not numeric_cells or not day_cells:
                continue
            rows.append((code, cells[1], parse_money_eur(numeric_cells[0]), int(day_cells[-1])))
    table_rows = normalize_ken_dataframe(pd.DataFrame(rows, columns=["ken_code", "ken_description", "basic_cost", "mean_duration_days"]))
    if not table_rows.empty:
        return table_rows
    return parse_ken_from_lines([paragraph.text for paragraph in document.paragraphs])


def load_ken_reference(source_dir: Path, out_ref_dir: Path):
    """Prefer improved official KEN CSVs, then parse official Word sources."""

    csv_candidates = [
        out_ref_dir / "ken.csv",
        source_dir / "data" / "reference" / "ken.csv",
        source_dir / "reference" / "ken.csv",
    ]
    if source_dir.exists():
        csv_candidates.extend(source_dir.rglob("ken.csv"))

    official_ken = []
    seen = set()
    for csv_path in csv_candidates:
        if not csv_path.exists() or csv_path in seen:
            continue
        seen.add(csv_path)
        try:
            candidate = normalize_ken_dataframe(pd.read_csv(csv_path, header=None))
        except Exception:
            try:
                candidate = normalize_ken_dataframe(pd.read_csv(csv_path))
            except Exception:
                continue
        if candidate.empty:
            continue
        if not is_placeholder_ken(candidate):
            official_ken.append(candidate)
    if official_ken:
        return max(official_ken, key=len), "official_csv"

    try:
        ken_path = resolve_source_file(source_dir, [
            "4.1 Λίστα Κλειστών Ενοποιημένων Νοσηλίων (2).docx",
            "4.1 Λίστα Κλειστών Ενοποιημένων Νοσηλίων (2).doc",
            "4.1 Λίστα Κλειστών Ενοποιημένων Νοσηλίων (3).docx",
            "4.1 Λίστα Κλειστών Ενοποιημένων Νοσηλίων (3).doc",
        ])
        source, kind = read_word_table_or_text(ken_path)
        ken = parse_ken_from_docx(source) if kind == "docx" else normalize_ken_dataframe(parse_ken_from_text(source))
        if not ken.empty and not is_placeholder_ken(ken):
            return ken, "official_word"
    except Exception:
        pass

    raise FileNotFoundError("Official KEN data was not found or could not be parsed. Provide the assignment KEN Word file or a cleaned official ken.csv.")


def build_generated_icd10_ken_map(icd: pd.DataFrame, ken: pd.DataFrame) -> pd.DataFrame:
    icd_prefixes = (
        icd["icd10_code"]
        .astype(str)
        .str.replace(".", "", regex=False)
        .str.extract(r"^([A-ZΑ-Ω]\d{2})")[0]
        .dropna()
        .drop_duplicates()
        .sort_values()
        .tolist()
    )
    preferred_prefixes = [
        "I21", "I25", "I50", "I48", "I20", "K35", "K40", "K45", "S72", "T81",
        "J96", "A41", "I46", "R57", "G93", "R07", "R10", "S09", "T14", "R55",
        "G40", "G41", "I63", "I64", "G45", "M16", "M17", "S82", "M25",
        "J18", "J44", "J45", "J20", "K52", "K80", "K92", "K29",
        "C34", "C18", "C50", "C71", "C25", "A09", "R50", "J21", "H66",
        "H25", "H26", "H33", "H40", "H10", "J34", "J02", "R04", "H81",
        "N18", "N17", "N20", "I12", "E11", "F32", "F20", "F41", "F10", "R45",
        "I10", "N39"
    ]
    available = set(icd_prefixes)
    ordered_prefixes = [pfx for pfx in preferred_prefixes if pfx in available]
    ordered_prefixes += [pfx for pfx in icd_prefixes if pfx not in set(ordered_prefixes)]
    ken_codes = ken["ken_code"].tolist()
    return pd.DataFrame([
        {"mdc_code": str((i % 25) + 1), "ken_code": ken_codes[i % len(ken_codes)], "icd10_code_prefix": pfx}
        for i, pfx in enumerate(ordered_prefixes)
    ])


def load_icd10_ken_map(source_dir: Path, icd: pd.DataFrame, ken: pd.DataFrame):
    """Use the official ICD10-KEN map when it is compatible with the active KEN table."""

    ken_codes = set(ken["ken_code"])
    icd_prefixes = set(
        icd["icd10_code"]
        .astype(str)
        .str.replace(".", "", regex=False)
        .str.extract(r"^([A-ZΑ-Ω]\d{2})")[0]
        .dropna()
    )
    try:
        map_path = resolve_source_file(source_dir, [
            "4.4 Λίστα Αντιστοιχήσεων ICD 10 - KEN (1).xlsx",
            "4.4 Λίστα Αντιστοιχήσεων ICD 10 - KEN (1).xls",
        ])
        raw = read_excel_flexible(map_path, sheet_name=0, header=0)
        if raw.shape[1] >= 3:
            official = pd.DataFrame({
                "mdc_code": raw.iloc[:, 0].map(lambda v: f"{int(v):02d}" if pd.notna(v) and str(v).replace(".", "", 1).isdigit() else normalize_space(v)),
                "ken_code": raw.iloc[:, 1].map(normalize_space).str.upper(),
                "icd10_code_prefix": raw.iloc[:, 2].map(normalize_space).str.upper().str.replace(".", "", regex=False),
            })
            official = official[
                official["ken_code"].isin(ken_codes)
                & official["icd10_code_prefix"].isin(icd_prefixes)
            ].drop_duplicates(["ken_code", "icd10_code_prefix"]).reset_index(drop=True)
            if not official.empty:
                return official, "official"
    except Exception:
        pass

    return build_generated_icd10_ken_map(icd, ken), "generated_compatible"


def load_reference_data(source_dir: Path, out_ref_dir: Path, ema_xlsx: Path | None = None):
    """Load official reference sources and write normalized reference CSVs."""

    out_ref_dir.mkdir(parents=True, exist_ok=True)
    existing_ref_dir = source_dir / "data" / "reference"
    existing_reference_files = {
        "icd10": "icd10_diagnosis.csv",
        "ken": "ken.csv",
        "icd10_ken": "icd10_ken_map.csv",
        "procedure_catalog": "procedure_catalog.csv",
        "lab_test_catalog": "lab_test_catalog.csv",
        "drug": "drug.csv",
        "active_substance": "active_substance.csv",
        "drug_active_substance": "drug_active_substance.csv",
    }
    if ema_xlsx is None and all((existing_ref_dir / name).exists() for name in existing_reference_files.values()):
        # Final-submission mode: when the repository already contains the
        # cleaned official reference CSVs, reuse them directly. This lets a
        # teammate regenerate only the synthetic transactional data without
        # needing the original Excel/Word workbooks on their laptop.
        frames = {
            key: pd.read_csv(existing_ref_dir / filename, dtype=str if key in {"icd10", "ken", "icd10_ken", "procedure_catalog", "lab_test_catalog", "drug"} else None)
            for key, filename in existing_reference_files.items()
        }
        for filename in existing_reference_files.values():
            src = existing_ref_dir / filename
            dst = out_ref_dir / filename
            if src.resolve() != dst.resolve():
                shutil.copy2(src, dst)
        metadata_src = existing_ref_dir / "reference_metadata.json"
        if metadata_src.exists() and metadata_src.resolve() != (out_ref_dir / "reference_metadata.json").resolve():
            shutil.copy2(metadata_src, out_ref_dir / "reference_metadata.json")
        meta = {"mode": "existing_clean_reference_csv", "ema_mode": "existing_clean_reference_csv"}
        metadata_src = existing_ref_dir / "reference_metadata.json"
        if metadata_src.exists():
            try:
                meta.update(json.loads(metadata_src.read_text(encoding="utf-8")))
            except json.JSONDecodeError:
                pass
        return {
            **frames,
            "meta": meta,
        }

    # ICD-10
    icd_path = resolve_source_file(source_dir, [
        "4.2 Κωδικοί ICD-10 15-12-2011 (2).xlsx",
        "4.2 Κωδικοί ICD-10 15-12-2011 (2).xls",
        "4.2 Κωδικοί ICD-10 15-12-2011 (2).xlsx",
        "4.2 Κωδικοί ICD-10 15-12-2011 (2).xls",
    ])
    icd_raw = read_excel_flexible(icd_path, sheet_name=0, header=None, names=["icd10_code", "icd10_description"])
    icd = icd_raw.copy()
    icd["icd10_code"] = icd["icd10_code"].map(normalize_space).str.upper()
    icd["icd10_description"] = icd["icd10_description"].map(normalize_space)
    icd = icd[icd["icd10_code"].notna() & icd["icd10_description"].notna()].drop_duplicates().reset_index(drop=True)
    icd.to_csv(out_ref_dir / "icd10_diagnosis.csv", index=False)

    # KEN: prefer the improved official CSV that may already exist in the
    # output bundle, then parse official Word sources.
    ken, ken_mode = load_ken_reference(source_dir, out_ref_dir)
    ken.to_csv(out_ref_dir / "ken.csv", index=False)

    # ICD-10 ↔ KEN: official mapping is used only after filtering it against
    # the exact KEN table being written, so hospitalizations cannot drift.
    icd_ken, icd_ken_mode = load_icd10_ken_map(source_dir, icd, ken)
    icd_ken.to_csv(out_ref_dir / "icd10_ken_map.csv", index=False)

    # Procedure catalog
    proc_path = resolve_source_file(source_dir, [
        "ΕΛΛΗΝΙΚΗ ΟΝΟΜΑΤΟΛΟΓΙΑ ΚΑΙ ΚΩΔΙΚΟΠΟΙΗΣΗ ΤΩΝ ΙΑΤΡΙΚΩΝ ΠΡΑΞΕΩΝ (2).xlsx",
        "ΕΛΛΗΝΙΚΗ ΟΝΟΜΑΤΟΛΟΓΙΑ ΚΑΙ ΚΩΔΙΚΟΠΟΙΗΣΗ ΤΩΝ ΙΑΤΡΙΚΩΝ ΠΡΑΞΕΩΝ (2).xls",
        "ΕΛΛΗΝΙΚΗ ΟΝΟΜΑΤΟΛΟΓΙΑ ΚΑΙ ΚΩΔΙΚΟΠΟΙΗΣΗ ΤΩΝ ΙΑΤΡΙΚΩΝ ΠΡΑΞΕΩΝ (3).xls",
        "ΕΛΛΗΝΙΚΗ ΟΝΟΜΑΤΟΛΟΓΙΑ ΚΑΙ ΚΩΔΙΚΟΠΟΙΗΣΗ ΤΩΝ ΙΑΤΡΙΚΩΝ ΠΡΑΞΕΩΝ (4).xls",
    ])
    proc_raw = read_excel_flexible(proc_path, sheet_name=0, header=1)
    proc = proc_raw.rename(columns={"ΚΩΔΙΚΟΣ": "procedure_code", "ΤΕΛΙΚΕΣ ΟΝΟΜΑΣΙΕΣ ΙΑΤΡΙΚΩΝ ΠΡΑΞΕΩΝ": "procedure_name"})[["procedure_code", "procedure_name"]].copy()
    proc["procedure_code"] = proc["procedure_code"].map(normalize_space)
    proc["procedure_name"] = proc["procedure_name"].map(normalize_space)
    proc = proc[proc["procedure_code"].notna() & proc["procedure_name"].notna()]
    proc = proc[proc["procedure_code"].str.match(r"^[A-ZΑ-Ω][A-Z0-9]{4,}$", na=False)].drop_duplicates("procedure_code").reset_index(drop=True)
    proc["procedure_category"] = proc["procedure_name"].apply(derive_proc_category)
    proc["standard_duration_min"] = proc.apply(lambda r: derive_standard_duration_min(r["procedure_name"], r["procedure_category"]), axis=1)
    proc["standard_cost"] = proc.apply(lambda r: derive_standard_cost(r["procedure_name"], r["procedure_category"]), axis=1)
    proc["required_place_type"] = proc.apply(lambda r: derive_place_type(r["procedure_name"], r["procedure_category"]), axis=1)
    proc = proc[[
        "procedure_code",
        "procedure_name",
        "procedure_category",
        "required_place_type",
        "standard_duration_min",
        "standard_cost",
    ]]
    proc[["procedure_code", "procedure_name", "procedure_category", "required_place_type"]].to_csv(
        out_ref_dir / "procedure_catalog.csv",
        index=False,
    )

    # Lab test catalog derived from official procedure file
    lab = proc[proc["procedure_category"] == "DIAGNOSTIC"].copy()
    lab["test_type"] = lab["procedure_name"].apply(derive_test_type)
    lab = lab.rename(columns={"procedure_code": "test_code", "procedure_name": "test_name"})[["test_code", "test_name", "test_type"]].drop_duplicates("test_code")
    lab.to_csv(out_ref_dir / "lab_test_catalog.csv", index=False)

    meta = {
        "seed": SEED,
        "note": "ICD-10, KEN, ICD10-KEN, and procedure references are cleaned from the best available official/improved sources. Demo reference rows are not generated for the final dataset.",
        "ken_mode": ken_mode,
        "icd10_ken_mode": icd_ken_mode,
        "ema_mode": None,
    }

    # Drugs / active substances
    if ema_xlsx and ema_xlsx.exists():
        # EMA Article 57 exports include title/metadata rows before the real
        # table header. Detect the header row instead of assuming row 1, so the
        # official workbook can be used directly.
        ema_probe = pd.read_excel(ema_xlsx, header=None, nrows=60)
        ema_header_row = None
        for idx, row in ema_probe.iterrows():
            labels = [str(v).strip().lower() for v in row.tolist() if pd.notna(v)]
            if any("product name" in v for v in labels) and any("active substance" in v for v in labels):
                ema_header_row = idx
                break
        if ema_header_row is None:
            raise ValueError("Could not locate the Article 57 product-data header row in EMA workbook.")
        ema = pd.read_excel(ema_xlsx, header=ema_header_row).dropna(how="all")
        # best-effort column detection
        cols_lower = {str(c).strip().lower(): c for c in ema.columns}
        drug_col = None
        substance_col = None
        for c in ema.columns:
            cl = str(c).lower()
            if drug_col is None and ("short name" in cl or "product name" in cl or "product short name" in cl):
                drug_col = c
            if substance_col is None and ("active substance" in cl):
                substance_col = c
        if drug_col is None or substance_col is None:
            raise ValueError("Could not locate product-name / active-substance columns in EMA Article 57 file.")
        d = ema[[drug_col, substance_col]].rename(columns={drug_col: "drug_name", substance_col: "active_substance_raw"}).copy()
        d["drug_name"] = d["drug_name"].map(normalize_space)
        d["active_substance_raw"] = d["active_substance_raw"].map(normalize_space)
        d = d[d["drug_name"].notna() & d["active_substance_raw"].notna()].drop_duplicates().reset_index(drop=True)
        d["drug_name"] = d["drug_name"].str.replace("\\", "/", regex=False).str.slice(0, 255)
        d["drug_id"] = ["EMA%07d" % (i + 1) for i in range(len(d))]
        drug = d[["drug_id", "drug_name"]].drop_duplicates("drug_id")
        substances = []
        edges = []
        substance_map = {}
        next_sid = 1
        for row in d.itertuples(index=False):
            parts = [normalize_space(x).replace("\\", "/")[:255] for x in re.split(r"\|", row.active_substance_raw)]
            for s in [p for p in parts if p]:
                if s not in substance_map:
                    substance_map[s] = next_sid
                    substances.append((next_sid, s))
                    next_sid += 1
                edges.append((row.drug_id, substance_map[s]))
        active_substance = pd.DataFrame(substances, columns=["substance_id", "substance_name"])
        drug_active_substance = pd.DataFrame(edges, columns=["drug_id", "substance_id"]).drop_duplicates()
        meta["ema_mode"] = "official"
    else:
        drug = pd.DataFrame(columns=["drug_id", "drug_name"])
        active_substance = pd.DataFrame(columns=["substance_id", "substance_name"])
        drug_active_substance = pd.DataFrame(columns=["drug_id", "substance_id"])
        meta["ema_mode"] = "missing"
        meta["note_drugs"] = "EMA Article 57 file was not provided. Drug, active-substance, allergy, and prescription CSVs are intentionally empty until an official EMA export is supplied."

    drug.to_csv(out_ref_dir / "drug.csv", index=False)
    active_substance.to_csv(out_ref_dir / "active_substance.csv", index=False)
    drug_active_substance.to_csv(out_ref_dir / "drug_active_substance.csv", index=False)

    (out_ref_dir / "reference_metadata.json").write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")
    return {
        "icd10": icd,
        "ken": ken,
        "icd10_ken": icd_ken,
        "procedure_catalog": proc,
        "lab_test_catalog": lab,
        "drug": drug,
        "active_substance": active_substance,
        "drug_active_substance": drug_active_substance,
        "meta": meta,
    }


def build_people_and_org(gen_dir: Path):
    """Create departments, staff, beds, and operating/procedure places."""

    departments = []
    floor_labels = ["B1/Building A", "1/Building A", "2/Building A", "3/Building A", "1/Building B", "2/Building B", "3/Building B"]
    for idx, (dname, _) in enumerate(DEPARTMENT_SPECS, start=1):
        bed_capacity = 18 + (idx % 5) * 4
        departments.append({
            "department_id": idx,
            "department_name": dname,
            "description": f"{dname} department of Ygeiopolis General Hospital",
            "bed_capacity": bed_capacity,
            "floor_building": floor_labels[idx % len(floor_labels)],
            "manager_doctor_amka": None,  # filled later
        })
    department_df = pd.DataFrame(departments)

    # Personnel pools
    personnel_rows, doctor_rows, nurse_rows, admin_rows, doctor_department_rows = [], [], [], [], []
    amka_counter = 10000000000
    person_idx = 1

    # Larger personnel pools make the three-year shift roster realistic:
    # every department must staff morning/afternoon/night shifts without
    # breaking monthly limits, rest windows, or senior-doctor coverage.
    for dep in DEPARTMENT_SPECS:
        dep_name, spec = dep
        dep_id = department_df.loc[department_df["department_name"] == dep_name, "department_id"].iloc[0]
        local_doctors = []
        ranks = (
            ["DIRECTOR"]
            + ["CONSULTANT_A"] * 8
            + ["CONSULTANT_B"] * 8
            + ["RESIDENT"] * 11
        )
        for rank in ranks:
            gender = pick(["M", "F"])
            first = pick(FIRST_NAMES_M if gender == "M" else FIRST_NAMES_F)
            last = pick(LAST_NAMES)
            amka = f"{amka_counter:011d}"; amka_counter += 1
            age = {
                "DIRECTOR": random.randint(50, 63),
                "CONSULTANT_A": random.randint(41, 54),
                "CONSULTANT_B": random.randint(33, 46),
                "RESIDENT": random.randint(27, 35),
            }[rank]
            personnel_rows.append({
                "amka": amka,
                "first_name": first,
                "last_name": last,
                "age": age,
                "email": mk_email(first, last, uniq=person_idx),
                "phone_number": phone(),
                "hiring_date": (date(2010,1,1) + timedelta(days=random.randint(0, 5500))).isoformat(),
                "personnel_type": "DOCTOR",
            })
            person_idx += 1
            local_doctors.append({"amka": amka, "rank": rank, "dep_id": dep_id})
        # supervisors within department, no cycles
        director_amka = [d["amka"] for d in local_doctors if d["rank"] == "DIRECTOR"][0]
        consultant_a_ambas = [d["amka"] for d in local_doctors if d["rank"] == "CONSULTANT_A"]
        consultant_b_ambas = [d["amka"] for d in local_doctors if d["rank"] == "CONSULTANT_B"]
        residents = [d["amka"] for d in local_doctors if d["rank"] == "RESIDENT"]
        department_df.loc[department_df["department_id"] == dep_id, "manager_doctor_amka"] = director_amka
        for d in local_doctors:
            if d["rank"] == "DIRECTOR":
                sup = None
            elif d["rank"] == "CONSULTANT_A":
                sup = director_amka
            elif d["rank"] == "CONSULTANT_B":
                sup = pick(consultant_a_ambas + [director_amka])
            else:
                sup = pick(consultant_a_ambas + consultant_b_ambas)
            doctor_rows.append({
                "amka": d["amka"],
                "license_number": f"LIC-{dep_id:02d}-{d['amka'][-5:]}",
                "specialization": spec,
                "doctor_rank": d["rank"],
                "supervisor_amka": sup,
            })
            doctor_department_rows.append({"doctor_amka": d["amka"], "department_id": dep_id})
            # add a small share to second department for cross-membership
            if d["rank"] in ("CONSULTANT_A", "CONSULTANT_B") and random.random() < 0.20:
                other = random.choice([x for x in department_df["department_id"] if x != dep_id])
                doctor_department_rows.append({"doctor_amka": d["amka"], "department_id": int(other)})

    # Nurses: 36 per department. This keeps enough capacity for dense,
    # multi-year shift assignments while preserving department ownership.
    for dep_id in department_df["department_id"]:
        for i in range(36):
            gender = pick(["M", "F"])
            first = pick(FIRST_NAMES_M if gender == "M" else FIRST_NAMES_F)
            last = pick(LAST_NAMES)
            amka = f"{amka_counter:011d}"; amka_counter += 1
            rank = "HEAD_NURSE" if i == 0 else ("ASSISTANT_NURSE" if i < 6 else "NURSE")
            age = {"HEAD_NURSE": random.randint(45, 58), "NURSE": random.randint(28, 50), "ASSISTANT_NURSE": random.randint(22, 35)}[rank]
            personnel_rows.append({
                "amka": amka,
                "first_name": first,
                "last_name": last,
                "age": age,
                "email": mk_email(first, last, uniq=person_idx),
                "phone_number": phone(),
                "hiring_date": (date(2012,1,1) + timedelta(days=random.randint(0, 5000))).isoformat(),
                "personnel_type": "NURSE",
            })
            person_idx += 1
            nurse_rows.append({
                "amka": amka,
                "nurse_rank": rank,
                "department_id": int(dep_id),
            })

    # Administrative staff: 12 per department, enough for continuous
    # admissions/front-desk coverage in the generated roster.
    for dep_id in department_df["department_id"]:
        for i in range(12):
            gender = pick(["M", "F"])
            first = pick(FIRST_NAMES_M if gender == "M" else FIRST_NAMES_F)
            last = pick(LAST_NAMES)
            amka = f"{amka_counter:011d}"; amka_counter += 1
            personnel_rows.append({
                "amka": amka,
                "first_name": first,
                "last_name": last,
                "age": random.randint(24, 58),
                "email": mk_email(first, last, uniq=person_idx),
                "phone_number": phone(),
                "hiring_date": (date(2013,1,1) + timedelta(days=random.randint(0, 4500))).isoformat(),
                "personnel_type": "ADMIN",
            })
            person_idx += 1
            admin_rows.append({
                "amka": amka,
                "admin_role": ADMIN_ROLES[i % len(ADMIN_ROLES)],
                "office_work": OFFICE_WORKS[i % len(OFFICE_WORKS)],
                "department_id": int(dep_id),
            })

    personnel_df = pd.DataFrame(personnel_rows)
    doctor_df = pd.DataFrame(doctor_rows)
    nurse_df = pd.DataFrame(nurse_rows)
    admin_df = pd.DataFrame(admin_rows)
    doctor_department_df = pd.DataFrame(doctor_department_rows).drop_duplicates()

    # Beds
    beds = []
    bed_id = 1
    for row in department_df.itertuples(index=False):
        for bed_number in range(1, int(row.bed_capacity) + 1):
            if row.department_name == "ICU" or bed_number <= 4:
                bed_type = "ICU"
            elif bed_number % 4 == 0:
                bed_type = "SINGLE"
            else:
                bed_type = "MULTI_BED"
            beds.append({
                "bed_id": bed_id,
                "department_id": row.department_id,
                "bed_number": bed_number,
                "bed_type": bed_type,
                "bed_status": pick(BED_STATUSES),
            })
            bed_id += 1
    bed_df = pd.DataFrame(beds)

    # Operating places
    operating_places = []
    for i, place_type in enumerate(PLACE_TYPES, start=1):
        operating_places.append({
            "place_id": i,
            "place_name": f"{'OR' if place_type == 'OPERATING_ROOM' else 'PR'}-{i:02d}",
            "place_type": place_type,
            "place_status": "AVAILABLE" if i % 5 else "MAINTENANCE",
        })
    operating_place_df = pd.DataFrame(operating_places)

    output_frames = {
        "department": department_df,
        "personnel": personnel_df,
        "doctor": doctor_df,
        "doctor_department": doctor_department_df,
        "nurse": nurse_df,
        "administrative_staff": admin_df,
        "bed": bed_df[["bed_id", "department_id", "bed_type", "bed_status"]],
        "operating_place": operating_place_df,
    }
    for name, df in output_frames.items():
        df.to_csv(gen_dir / f"{name}.csv", index=False)

    return {
        "department": department_df,
        "personnel": personnel_df,
        "doctor": doctor_df,
        "doctor_department": doctor_department_df,
        "nurse": nurse_df,
        "administrative_staff": admin_df,
        "bed": bed_df,
        "operating_place": operating_place_df,
    }


def build_patients(gen_dir: Path, count=500):
    """Create patients and optional emergency contacts."""

    patient_rows, contact_rows = [], []
    amka_counter = 30000000000
    for i in range(count):
        gender = pick(["M", "F"])
        first = pick(FIRST_NAMES_M if gender == "M" else FIRST_NAMES_F)
        last = pick(LAST_NAMES)
        father = pick(FATHER_NAMES)
        amka = f"{amka_counter:011d}"; amka_counter += 1
        age = random.randint(1, 92)
        patient_rows.append({
            "patient_amka": amka,
            "first_name": first,
            "last_name": last,
            "father_name": father,
            "age": age,
            "gender": "MALE" if gender == "M" else "FEMALE",
            "weight_kg": round(random.uniform(8, 120), 2) if age > 1 else round(random.uniform(3, 12), 2),
            "height_cm": round(random.uniform(65, 195), 2),
            "address_line": f"{pick(['Athens','Piraeus','Patras','Larisa','Volos'])} {random.randint(1,150)}",
            "phone_number": phone(),
            "email": mk_email(first, last, domain="mail.example", uniq=i+1),
            "profession": pick(PROFESSIONS),
            "nationality": pick(NATIONALITIES),
            "insurance_provider": random.choices(INSURANCE_PROVIDERS, weights=[55, 15, 10, 10, 10])[0],
        })
        for _ in range(random.randint(0, 2)):
            cf = pick(FIRST_NAMES_M + FIRST_NAMES_F)
            cl = pick(LAST_NAMES)
            contact_rows.append({
                "patient_amka": amka,
                "first_name": cf,
                "last_name": cl,
                "phone_number": phone(prefix="21"),
                "email": mk_email(cf, cl, domain="contact.example", uniq=random.randint(1,9999)),
            })
    patient_df = pd.DataFrame(patient_rows)
    emergency_contact_df = pd.DataFrame(contact_rows).drop_duplicates()
    patient_df.to_csv(gen_dir / "patient.csv", index=False)
    emergency_contact_df.to_csv(gen_dir / "emergency_contact.csv", index=False)
    return {"patient": patient_df, "emergency_contact": emergency_contact_df}


def build_shifts(
    gen_dir: Path,
    org,
    start_date=DATASET_START_DATE,
    days=DEFAULT_SHIFT_DAYS,
    sample_days_per_month=DEFAULT_SHIFT_SAMPLE_DAYS_PER_MONTH,
):
    """Create a multi-year roster sample that respects the schema shift rules."""

    dept_df = org["department"]
    doctor_df = org["doctor"]
    nurse_df = org["nurse"]
    admin_df = org["administrative_staff"]
    shifts, assignments = [], []
    shift_id = 1
    monthly_counts = defaultdict(int)
    last_shift_end = {}
    consecutive_nights = defaultdict(int)

    def shift_bounds(shift_date, shift_type):
        start_text, end_text = SHIFT_TIMES[shift_type]
        start_dt = datetime.combine(shift_date, time.fromisoformat(start_text))
        end_dt = datetime.combine(shift_date, time.fromisoformat(end_text))
        if shift_type == "NIGHT":
            end_dt += timedelta(days=1)
        return start_text, end_text, start_dt, end_dt

    def eligible(amka, start_dt, shift_type, monthly_limit):
        month_key = (amka, start_dt.year, start_dt.month)
        if monthly_counts[month_key] >= monthly_limit:
            return False
        previous_end = last_shift_end.get(amka)
        if previous_end is not None and start_dt - previous_end < timedelta(hours=8):
            return False
        if shift_type == "NIGHT" and consecutive_nights[amka] >= 3:
            return False
        return True

    def register(amka, start_dt, end_dt, shift_type):
        monthly_counts[(amka, start_dt.year, start_dt.month)] += 1
        last_shift_end[amka] = end_dt
        if shift_type == "NIGHT":
            consecutive_nights[amka] += 1
        else:
            consecutive_nights[amka] = 0

    def pick_available(pool, needed, start_dt, end_dt, shift_type, monthly_limit, selected=None):
        selected = set(selected or [])
        candidates = [
            amka for amka in pool
            if amka not in selected and eligible(amka, start_dt, shift_type, monthly_limit)
        ]
        random.shuffle(candidates)
        candidates.sort(key=lambda amka: monthly_counts[(amka, start_dt.year, start_dt.month)])
        chosen = candidates[:needed]
        if len(chosen) < needed:
            raise ValueError(
                "Could not build a valid shift roster. Increase department staff pools "
                "or reduce --shift-days."
            )
        for amka in chosen:
            register(amka, start_dt, end_dt, shift_type)
        return chosen

    def shift_dates_for_window():
        all_dates = date_range(start_date, days)
        if sample_days_per_month <= 0:
            return all_dates
        by_month = defaultdict(list)
        for shift_date in all_dates:
            by_month[(shift_date.year, shift_date.month)].append(shift_date)
        selected_dates = []
        for month_key in sorted(by_month):
            # A full week per month keeps the roster visibly spread across
            # three years without making LOAD DATA impractical under triggers.
            selected_dates.extend(by_month[month_key][:sample_days_per_month])
        return selected_dates

    # pools by department
    doctors_by_dep = defaultdict(list)
    for d in doctor_df.itertuples(index=False):
        dep_id = org["doctor_department"].loc[org["doctor_department"]["doctor_amka"] == d.amka, "department_id"].iloc[0]
        doctors_by_dep[int(dep_id)].append({"amka": d.amka, "rank": d.doctor_rank})
    nurses_by_dep = {int(dep): list(g["amka"]) for dep, g in nurse_df.groupby("department_id")}
    admins_by_dep = {int(dep): list(g["amka"]) for dep, g in admin_df.groupby("department_id")}

    for dep in dept_df.itertuples(index=False):
        dep_id = int(dep.department_id)
        dep_docs = doctors_by_dep[dep_id]
        seniors = [x["amka"] for x in dep_docs if x["rank"] in ("DIRECTOR", "CONSULTANT_A")]
        others = [x["amka"] for x in dep_docs]
        nurse_pool = nurses_by_dep[dep_id]
        admin_pool = admins_by_dep[dep_id]

        for d in shift_dates_for_window():
            for stype_idx, stype in enumerate(["MORNING","AFTERNOON","NIGHT"]):
                start_time, end_time, start_dt, end_dt = shift_bounds(d, stype)
                shifts.append({
                    "shift_id": shift_id,
                    "department_id": dep_id,
                    "shift_date": d.isoformat(),
                    "shift_type": stype,
                    "start_time": start_time,
                    "end_time": end_time,
                    "shift_status": "PROCESSING",
                })
                # Shift coverage rule: every shift has three doctors, including
                # at least one senior doctor so residents are never unsupervised.
                doc_candidates = pick_available(seniors, 1, start_dt, end_dt, stype, 15)
                doc_candidates += pick_available(others, 2, start_dt, end_dt, stype, 15, selected=doc_candidates)
                for amka in doc_candidates:
                    assignments.append({"shift_id": shift_id, "personnel_amka": amka, "assigned_role": "ON_CALL_DOCTOR"})

                # Nursing/admin coverage follows the project business rules:
                # six nurses and two administrative staff per department shift.
                nurse_candidates = pick_available(nurse_pool, 6, start_dt, end_dt, stype, 20)
                for amka in nurse_candidates:
                    assignments.append({"shift_id": shift_id, "personnel_amka": amka, "assigned_role": "ON_CALL_NURSE"})

                admin_candidates = pick_available(admin_pool, 2, start_dt, end_dt, stype, 25)
                for amka in admin_candidates:
                    assignments.append({"shift_id": shift_id, "personnel_amka": amka, "assigned_role": "ON_CALL_ADMIN"})

                shift_id += 1

    shift_df = pd.DataFrame(shifts)
    assign_df = pd.DataFrame(assignments).drop_duplicates()
    shift_df.to_csv(gen_dir / "department_shift.csv", index=False)
    assign_df.to_csv(gen_dir / "shift_assignment.csv", index=False)
    return {"department_shift": shift_df, "shift_assignment": assign_df}


def build_emergency_visits(gen_dir: Path, patients, org, count=18000, start_date=DATASET_START_DATE, days=DEFAULT_SHIFT_DAYS):
    """Create emergency visits with triage levels and service timestamps."""

    patient_df = patients["patient"]
    nurse_df = org["nurse"]
    dep_df = org["department"]

    emergency_dep_id = int(dep_df.loc[dep_df["department_name"] == "Emergency", "department_id"].iloc[0])
    non_emergency_depts = [int(x) for x in dep_df["department_id"] if x != emergency_dep_id]
    triage_nurses = nurse_df.loc[nurse_df["department_id"] == emergency_dep_id, "amka"].tolist()
    if len(triage_nurses) < 3:
        triage_nurses = nurse_df["amka"].tolist()[:6]

    visits = []
    visit_id = 1
    start_dt = datetime.combine(start_date, time(0, 0, 0))
    for i in range(count):
        arrival = start_dt + timedelta(hours=random.randint(0, 24 * days - 1), minutes=random.randint(0, 59))
        level = random.choices([1,2,3,4,5], weights=[6,14,28,32,20])[0]
        wait_minutes = {
            1: random.randint(1, 8),
            2: random.randint(5, 20),
            3: random.randint(10, 60),
            4: random.randint(20, 120),
            5: random.randint(30, 180),
        }[level]
        is_waiting = random.random() < 0.15
        service_start = arrival + timedelta(minutes=wait_minutes)
        service_end = service_start + timedelta(minutes=random.randint(20, 180))
        hospitalized = random.random() < {1: 0.92, 2: 0.74, 3: 0.48, 4: 0.22, 5: 0.08}[level]
        referred_dep = pick(non_emergency_depts) if hospitalized else ""
        visits.append({
            "visit_id": visit_id,
            "patient_amka": pick(patient_df["patient_amka"]),
            "triage_nurse_amka": pick(triage_nurses),
            "arrival_ts": arrival.strftime("%Y-%m-%d %H:%M:%S"),
            "symptoms": pick(SYMPTOMS_BY_LEVEL[level]),
            "emergency_level": level,
            "service_start_ts": "" if is_waiting else service_start.strftime("%Y-%m-%d %H:%M:%S"),
            "disposition": "HOSPITALIZED" if hospitalized else "DISCHARGED",
            "referred_department_id": referred_dep,
            "discharge_instructions": "" if hospitalized else "Rest, hydration, and outpatient follow-up",
            "status": "WAITING" if is_waiting else "CALLED",
        })
        visit_id += 1

    ev_df = pd.DataFrame(visits)
    ev_df.to_csv(gen_dir / "emergency_visit.csv", index=False)
    return {"emergency_visit": ev_df}


def build_clinical_events(
    gen_dir: Path,
    ref,
    org,
    patients,
    hospitalization_count=12000,
    lab_test_count=12000,
    procedure_event_count=6000,
    prescription_count=12000,
    admission_start_date=DATASET_START_DATE,
    admission_days=DEFAULT_SHIFT_DAYS,
):
    """Create hospitalizations and related clinical events."""

    dept_df = org["department"]
    bed_df = org["bed"]
    patient_df = patients["patient"]
    doctor_df = org["doctor"]
    nurse_df = org["nurse"]
    dep_name_to_id = {r.department_name: int(r.department_id) for r in dept_df.itertuples(index=False)}

    # Prepare mappings
    icd_df = ref["icd10"].copy()
    icd_df["prefix3"] = icd_df["icd10_code"].str.replace(".", "", regex=False).str.extract(r"^([A-ZΑ-Ω]\d{2})")[0]
    prefix_to_codes = defaultdict(list)
    for row in icd_df.itertuples(index=False):
        prefix_to_codes[row.prefix3].append(row.icd10_code)

    valid_ken_codes = set(ref["ken"]["ken_code"])
    map_df = ref["icd10_ken"]
    map_df = map_df[map_df["ken_code"].isin(valid_ken_codes)].copy()
    if map_df.empty:
        raise ValueError("ICD10-KEN map has no rows compatible with ken.csv.")
    prefix_to_kens = defaultdict(list)
    for row in map_df.itertuples(index=False):
        prefix_to_kens[row.icd10_code_prefix].append(row.ken_code)

    # Prefix preferences per department
    pref = {
        "Cardiology": ["I21", "I25", "I50", "I48", "I20"],
        "Surgery": ["K35", "K40", "K45", "S72", "T81"],
        "ICU": ["J96", "A41", "I46", "R57", "G93"],
        "Emergency": ["R07", "R10", "S09", "T14", "R55"],
        "Neurology": ["G40", "G41", "I63", "I64", "G45"],
        "Orthopedics": ["M16", "M17", "S72", "S82", "M25"],
        "Pulmonology": ["J18", "J44", "J45", "J96", "J20"],
        "Gastroenterology": ["K35", "K52", "K80", "K92", "K29"],
        "Oncology": ["C34", "C18", "C50", "C71", "C25"],
        "Pediatrics": ["J18", "A09", "R50", "J21", "H66"],
        "Ophthalmology": ["H25", "H26", "H33", "H40", "H10"],
        "ENT": ["H66", "J34", "J02", "R04", "H81"],
        "Nephrology": ["N18", "N17", "N20", "I12", "E11"],
        "Psychiatry": ["F32", "F20", "F41", "F10", "R45"],
        "Internal Medicine": ["E11", "I10", "J18", "N39", "K52"],
    }

    # build quick lookup doctors by dept
    home_dep = {}
    for row in org["doctor_department"].itertuples(index=False):
        home_dep.setdefault(row.doctor_amka, row.department_id)
    docs_by_dep = defaultdict(list)
    for row in doctor_df.itertuples(index=False):
        depid = int(home_dep[row.amka])
        docs_by_dep[depid].append({"amka": row.amka, "rank": row.doctor_rank, "specialization": row.specialization})
    nurses_by_dep = {int(dep): list(g["amka"]) for dep, g in nurse_df.groupby("department_id")}

    allocator = BedAllocator(bed_df)

    hosp_rows = []
    hosp_doctor_rows = []
    eval_rows = []
    lab_rows = []
    proc_event_rows = []
    proc_part_rows = []
    allergy_rows = []
    prescription_rows = []
    image_rows = []
    entity_image_rows = []

    # precompute special patients for query patterns
    patient_ids = list(patient_df["patient_amka"])
    repeat_same_dept_patients = random.sample(patient_ids, 15)
    equal_total_pairs = [tuple(random.sample([p for p in patient_ids if p not in repeat_same_dept_patients], 2)) for _ in range(10)]
    q14_prefixes = [p for p in ["I21","J18","K35","M16","G40"] if p in prefix_to_kens][:5]

    # Base hospitalization targets by department. The default three-year
    # window creates enough fact rows for visible EXPLAIN ANALYZE differences
    # when comparing indexed and non-indexed query plans.
    dept_ids = list(dept_df["department_id"])
    clinical_start = datetime.combine(admission_start_date, time(8, 0, 0))

    def pick_hosp_prefix(dep_name, forced_prefix=None):
        if forced_prefix:
            return forced_prefix
        cands = [p for p in pref.get(dep_name, []) if p in prefix_to_kens and p in prefix_to_codes]
        if not cands:
            cands = [p for p in prefix_to_kens.keys() if p in prefix_to_codes]
        return pick(cands)

    def pick_ken_for_prefix(prefix):
        return pick(prefix_to_kens[prefix])

    def pick_full_icd(prefix):
        return pick(prefix_to_codes[prefix])

    def compute_total_cost(ken_code, admission_ts, discharge_ts):
        k = ken_row_lookup[ken_code]
        actual_days = max(1, math.ceil((discharge_ts - admission_ts).total_seconds() / 86400))
        extra_days = max(0, actual_days - int(k.mean_duration_days))
        total_cost = round(float(k.basic_cost) + extra_days * float(k.extra_daily_cost), 2)
        return actual_days, total_cost

    ken_row_lookup = {r.ken_code: r for r in ref["ken"].itertuples(index=False)}
    proc_cat = ref["procedure_catalog"].drop_duplicates("procedure_code")
    proc_place_type = dict(proc_cat[["procedure_code", "required_place_type"]].itertuples(index=False, name=None))
    surg_proc = proc_cat[proc_cat["procedure_category"] == "SURGICAL"]["procedure_code"].tolist()
    diag_proc = proc_cat[proc_cat["procedure_category"] == "DIAGNOSTIC"]["procedure_code"].tolist()
    if not surg_proc and not diag_proc:
        raise ValueError("procedure_catalog.csv has no usable procedure codes.")
    if not surg_proc:
        surg_proc = diag_proc
    if not diag_proc:
        diag_proc = surg_proc
    lab_cat = ref["lab_test_catalog"]

    # --- hospitalizations: targeted first ---
    hosp_id = 1
    used_hosp_patient = set()
    patient_schedules = defaultdict(list)

    def patient_is_available(patient_amka, start_ts, end_ts):
        return all(not (start_ts < e and s < end_ts) for s, e in patient_schedules[patient_amka])

    def add_hospitalization(patient_amka, department_id, admission_ts, stay_days, prefix=None, forced_ken=None):
        nonlocal hosp_id
        dep_name = dept_df.loc[dept_df["department_id"] == department_id, "department_name"].iloc[0]
        prefix_local = pick_hosp_prefix(dep_name, forced_prefix=prefix)
        ken_code = forced_ken or pick_ken_for_prefix(prefix_local)
        adm_code = pick_full_icd(prefix_local)
        discharge_code = adm_code if random.random() < 0.65 else pick_full_icd(prefix_local)
        discharge_ts = admission_ts + timedelta(days=stay_days, hours=random.randint(2, 18))
        if not patient_is_available(patient_amka, admission_ts, discharge_ts):
            return False
        bed_id = allocator.allocate(department_id, admission_ts, discharge_ts)
        if bed_id is None:
            return False
        actual_days, total_cost = compute_total_cost(ken_code, admission_ts, discharge_ts)
        hosp_rows.append({
            "hosp_id": hosp_id,
            "patient_amka": patient_amka,
            "department_id": department_id,
            "bed_id": bed_id,
            "ken_code": ken_code,
            "admission_ts": admission_ts.strftime("%Y-%m-%d %H:%M:%S"),
            "discharge_ts": discharge_ts.strftime("%Y-%m-%d %H:%M:%S"),
            "admission_icd10_code": adm_code,
            "discharge_icd10_code": discharge_code,
            "total_cost": total_cost,
        })
        # doctors
        dep_docs = docs_by_dep[int(department_id)]
        primary = pick([d["amka"] for d in dep_docs if d["rank"] in ("DIRECTOR","CONSULTANT_A","CONSULTANT_B")])
        hosp_doctor_rows.append({"hosp_id": hosp_id, "doctor_amka": primary})
        others = random.sample([d["amka"] for d in dep_docs if d["amka"] != primary], k=random.randint(0,2))
        for od in others:
            hosp_doctor_rows.append({"hosp_id": hosp_id, "doctor_amka": od})
        # evaluation on most discharges
        if random.random() < 0.74:
            eval_rows.append({
                "hosp_id": hosp_id,
                "evaluation_date": (discharge_ts.date() + timedelta(days=random.randint(1,20))).isoformat(),
                "medical_care_score": random.choices([2,3,4,5],[5,15,35,45])[0],
                "nursing_care_score": random.choices([2,3,4,5],[4,16,38,42])[0],
                "cleanliness_score": random.choices([2,3,4,5],[6,20,36,38])[0],
                "food_score": random.choices([1,2,3,4,5],[5,12,25,35,23])[0],
                "overall_experience_score": random.choices([2,3,4,5],[5,15,40,40])[0],
                "comments": pick([
                    "Smooth hospitalization and clear discharge instructions",
                    "Very good medical team, waiting time acceptable",
                    "Nursing staff was attentive and professional",
                    "Overall positive experience",
                    "Good treatment outcome and organized follow-up",
                ]),
            })
        used_hosp_patient.add((hosp_id, patient_amka))
        patient_schedules[patient_amka].append((admission_ts, discharge_ts))
        hosp_id += 1
        return True

    # Q3 pattern: >3 hospitalizations in same department
    for p in repeat_same_dept_patients:
        dep_id = pick(dept_ids)
        base = datetime(2025, random.randint(1, 12), random.randint(1, 20), 9, 0, 0)
        for j in range(random.randint(4, 5)):
            add_hospitalization(
                p, dep_id, base + timedelta(days=70*j + random.randint(0, 10)),
                stay_days=random.randint(2, 6),
            )

    # Q9 pattern: pairs with same total days within 2026 > 15
    for a, b in equal_total_pairs:
        total_days = pick([18, 20, 22, 24])
        splits = sorted(random.sample(range(2, total_days-1), 2))
        parts = [splits[0], splits[1] - splits[0], total_days - splits[1]]
        for pat in [a, b]:
            dep_choices = random.sample(dept_ids, 3)
            base = datetime(2026, random.randint(1, 6), random.randint(1, 18), 10, 0, 0)
            for k, part in enumerate(parts):
                add_hospitalization(pat, dep_choices[k], base + timedelta(days=45*k), stay_days=part)

    # Q14 pattern: same ICD-category counts in two consecutive years
    for prefix in q14_prefixes:
        for year in [2025, 2026]:
            for _ in range(6):
                p = pick(patient_ids)
                dep_id = pick(dept_ids)
                admission = datetime(year, random.randint(1, 11), random.randint(1, 20), 11, 0, 0)
                add_hospitalization(p, dep_id, admission, stay_days=random.randint(1, 5), prefix=prefix)

    # Fill to the requested target while preserving no-overlap bed and patient rules.
    attempts = 0
    max_attempts = hospitalization_count * 80
    while len(hosp_rows) < hospitalization_count and attempts < max_attempts:
        attempts += 1
        p = pick(patient_ids)
        dep_id = pick(dept_ids)
        admission = clinical_start + timedelta(days=random.randint(0, admission_days - 1), hours=random.randint(0, 16))
        prefix = None
        add_hospitalization(p, dep_id, admission, stay_days=random.randint(1, 12), prefix=prefix)

    if len(hosp_rows) < hospitalization_count:
        raise ValueError(f"Could only generate {len(hosp_rows)} hospitalizations without bed/patient overlaps.")

    hosp_df = pd.DataFrame(hosp_rows)
    hosp_doc_df = pd.DataFrame(hosp_doctor_rows).drop_duplicates()
    eval_df = pd.DataFrame(eval_rows).drop_duplicates("hosp_id")

    # Lab tests
    test_id = 1
    lab_codes = list(lab_cat["test_code"])
    for row in hosp_df.sample(n=min(lab_test_count, len(hosp_df)), random_state=SEED).itertuples(index=False):
        ordered_by = pick(hosp_doc_df.loc[hosp_doc_df["hosp_id"] == row.hosp_id, "doctor_amka"])
        start_dt = datetime.fromisoformat(row.admission_ts)
        end_dt = datetime.fromisoformat(row.discharge_ts)
        test_dt = start_dt + timedelta(hours=random.randint(4, max(5, int((end_dt-start_dt).total_seconds()//3600)-1)))
        result_numeric = round(random.uniform(0.5, 180.0), 2) if random.random() < 0.55 else ""
        result_unit = pick(["mg/dL", "g/L", "mmol/L", "IU/L", "cells/uL"]) if result_numeric != "" else ""
        lab_rows.append({
            "test_id": test_id,
            "hosp_id": row.hosp_id,
            "test_code": pick(lab_codes),
            "ordered_by_doctor_amka": ordered_by,
            "test_datetime": test_dt.strftime("%Y-%m-%d %H:%M:%S"),
            "result_text": pick(["Normal", "Mildly abnormal", "Follow-up required", "Improved from previous"]),
        })
        test_id += 1
    lab_df = pd.DataFrame(lab_rows)

    # Procedure events and participants, with skew for top surgeons
    proc_event_id = 1
    place_schedules = defaultdict(list)
    doctor_schedules = defaultdict(list)
    staff_schedules = defaultdict(list)

    def staff_is_available(amka, start_ts, end_ts):
        return all(not (start_ts < e and s < end_ts) for s, e in staff_schedules[amka])

    # pick top surgeons to dominate
    surgeon_candidates = doctor_df[doctor_df["specialization"].isin(["GENERAL_SURGERY","ORTHOPEDICS","CARDIOLOGY","GASTROENTEROLOGY","NEUROLOGY"])]
    top_surgeons = surgeon_candidates[surgeon_candidates["doctor_rank"].isin(["DIRECTOR","CONSULTANT_A","CONSULTANT_B"])].sample(n=12, random_state=SEED)["amka"].tolist()
    weighted_surgeons = top_surgeons + surgeon_candidates["amka"].tolist()

    surgical_hosps = hosp_df.sample(n=min(procedure_event_count, len(hosp_df)), random_state=SEED)
    for row in surgical_hosps.itertuples(index=False):
        proc_code = pick(surg_proc if random.random() < 0.82 else diag_proc)
        place_type = proc_place_type[proc_code]
        place_ids = org["operating_place"].loc[org["operating_place"]["place_type"] == place_type, "place_id"].tolist()
        if not place_ids:
            continue
        chief = random.choice(weighted_surgeons if random.random() < 0.65 else surgeon_candidates["amka"].tolist())
        start_window = datetime.fromisoformat(row.admission_ts) + timedelta(hours=8)
        end_window = datetime.fromisoformat(row.discharge_ts) - timedelta(hours=6)
        if end_window <= start_window:
            continue
        duration = random.randint(45, 240)
        placed = False
        for _ in range(40):
            place_id = pick(place_ids)
            start_ts = start_window + timedelta(hours=random.randint(0, max(1, int((end_window - start_window).total_seconds() // 3600) - 1)))
            end_ts = start_ts + timedelta(minutes=duration)
            if end_ts > end_window:
                continue
            # overlap checks
            if any(start_ts < e and s < end_ts for s, e in place_schedules[place_id]):
                continue
            if any(start_ts < e and s < end_ts for s, e in doctor_schedules[chief]):
                continue
            if not staff_is_available(chief, start_ts, end_ts):
                continue
            place_schedules[place_id].append((start_ts, end_ts))
            doctor_schedules[chief].append((start_ts, end_ts))
            staff_schedules[chief].append((start_ts, end_ts))
            proc_event_rows.append({
                "procedure_event_id": proc_event_id,
                "hosp_id": row.hosp_id,
                "procedure_code": proc_code,
                "place_id": place_id,
                "chief_surgeon_amka": chief,
                "start_ts": start_ts.strftime("%Y-%m-%d %H:%M:%S"),
                "end_ts": end_ts.strftime("%Y-%m-%d %H:%M:%S"),
                "actual_duration_min": duration,
            })
            # participants: 1-2 doctors + 1 nurse
            dep_id = int(row.department_id)
            dep_doc_pool = [
                d["amka"]
                for d in docs_by_dep[dep_id]
                if d["amka"] != chief and staff_is_available(d["amka"], start_ts, end_ts)
            ]
            dep_nurse_pool = [
                amka for amka in nurses_by_dep[dep_id]
                if staff_is_available(amka, start_ts, end_ts)
            ]
            for helper in random.sample(dep_doc_pool, k=min(len(dep_doc_pool), random.randint(1, 2))):
                proc_part_rows.append({"procedure_event_id": proc_event_id, "personnel_amka": helper})
                staff_schedules[helper].append((start_ts, end_ts))
            if dep_nurse_pool:
                nurse_helper = pick(dep_nurse_pool)
                proc_part_rows.append({"procedure_event_id": proc_event_id, "personnel_amka": nurse_helper})
                staff_schedules[nurse_helper].append((start_ts, end_ts))
            proc_event_id += 1
            placed = True
            break
        if not placed:
            continue

    proc_event_df = pd.DataFrame(proc_event_rows)
    proc_part_df = pd.DataFrame(proc_part_rows).drop_duplicates()

    # Keep the loaded bed table consistent with the dataset validation date.
    # Historic hospitalizations do not change current bed status, but active
    # hospitalizations at DATASET_AS_OF_TS should be reflected as occupied.
    active_bed_ids = set(
        hosp_df.loc[
            (pd.to_datetime(hosp_df["admission_ts"]) <= DATASET_AS_OF_TS)
            & (pd.to_datetime(hosp_df["discharge_ts"]) > DATASET_AS_OF_TS),
            "bed_id",
        ]
    )
    bed_status_df = org["bed"].copy()
    bed_status_df.loc[bed_status_df["bed_id"].isin(active_bed_ids), "bed_status"] = "OCCUPIED"
    bed_status_df.loc[
        (~bed_status_df["bed_id"].isin(active_bed_ids)) & (bed_status_df["bed_status"] == "OCCUPIED"),
        "bed_status",
    ] = "AVAILABLE"
    bed_status_df[["bed_id", "department_id", "bed_type", "bed_status"]].to_csv(gen_dir / "bed.csv", index=False)
    org["bed"] = bed_status_df

    allergy_columns = ["patient_amka", "substance_id"]
    prescription_columns = [
        "prescription_id",
        "hosp_id",
        "patient_amka",
        "doctor_amka",
        "drug_id",
        "dosage",
        "frequency",
        "start_datetime",
        "end_datetime",
    ]

    # Allergies and prescriptions are generated only when official EMA-derived
    # drug/substance references exist. Without EMA, these CSVs stay empty
    # instead of introducing fake medication data into the final database.
    substances = ref["active_substance"].copy()
    substance_ids = list(substances["substance_id"])
    drug_df = ref["drug"].copy()
    das = ref["drug_active_substance"].copy()
    if not substance_ids or drug_df.empty or das.empty:
        allergy_df = pd.DataFrame(columns=allergy_columns)
        presc_df = pd.DataFrame(columns=prescription_columns)
    else:
        common_substance_ids = {s: int(substances.loc[substances["substance_name"] == s, "substance_id"].iloc[0]) for pair in COMMON_SUBSTANCE_PAIRS for s in pair if s in set(substances["substance_name"])}
        allergy_patients = random.sample(patient_ids, 70)
        for p in allergy_patients:
            chosen = set(random.sample(substance_ids, k=random.randint(1, 3)))
            # make popular allergy substances appear more often
            if common_substance_ids and random.random() < 0.45:
                chosen.add(pick(list(common_substance_ids.values())))
            for sid in chosen:
                allergy_rows.append({"patient_amka": p, "substance_id": sid})
        allergy_df = pd.DataFrame(allergy_rows).drop_duplicates()

        # Prescriptions (query-driven for Q10)
        sub_name = dict(ref["active_substance"][["substance_id","substance_name"]].itertuples(index=False, name=None))
        drug_to_subs = defaultdict(set)
        for r in das.itertuples(index=False):
            drug_to_subs[r.drug_id].add(sub_name[r.substance_id])
        all_drug_ids = list(drug_df["drug_id"])
        patient_banned_substances = {
            patient_amka: {sub_name[sid] for sid in group["substance_id"].tolist()}
            for patient_amka, group in allergy_df.groupby("patient_amka")
        }
        safe_drug_cache = {}

        def safe_drugs_for_patient(patient_amka):
            if patient_amka in safe_drug_cache:
                return safe_drug_cache[patient_amka]
            banned = patient_banned_substances.get(patient_amka, set())
            if not banned:
                safe_drug_cache[patient_amka] = all_drug_ids
            else:
                safe_drug_cache[patient_amka] = [drug_id for drug_id in all_drug_ids if not (drug_to_subs[drug_id] & banned)]
            return safe_drug_cache[patient_amka]

        # helper: pick a drug containing a target substance
        drugs_by_sub = defaultdict(list)
        for d, subs in drug_to_subs.items():
            for s in subs:
                drugs_by_sub[s].append(d)

        prescription_id = 1
        # target pairs
        hosp_for_pairs = hosp_df.sample(n=min(90, len(hosp_df)), random_state=SEED)
        for i, row in enumerate(hosp_for_pairs.itertuples(index=False)):
            pair = COMMON_SUBSTANCE_PAIRS[i % len(COMMON_SUBSTANCE_PAIRS)]
            safe = set(safe_drugs_for_patient(row.patient_amka))
            d1_candidates = [d for d in drugs_by_sub.get(pair[0], []) if d in safe]
            d2_candidates = [d for d in drugs_by_sub.get(pair[1], []) if d in safe and d not in d1_candidates]
            if not d1_candidates or not d2_candidates:
                continue
            admission_ts = datetime.fromisoformat(row.admission_ts)
            discharge_ts = datetime.fromisoformat(row.discharge_ts)
            max_start_hours = max(1, int((discharge_ts - admission_ts).total_seconds() // 3600) - 2)
            if max_start_hours < 3:
                continue
            start_base = admission_ts + timedelta(hours=random.randint(3, min(24, max_start_hours)))
            end_dt = min(discharge_ts, start_base + timedelta(days=random.randint(2, 6)))
            if end_dt <= start_base:
                continue
            doctor_choices = hosp_doc_df.loc[hosp_doc_df["hosp_id"] == row.hosp_id, "doctor_amka"].tolist()
            for drug_id in [pick(d1_candidates), pick(d2_candidates)]:
                prescription_rows.append({
                    "prescription_id": prescription_id,
                    "hosp_id": row.hosp_id,
                    "patient_amka": row.patient_amka,
                    "doctor_amka": pick(doctor_choices),
                    "drug_id": drug_id,
                    "dosage": pick(["1 tablet", "500 mg", "1 vial", "1 capsule", "40 mg"]),
                    "frequency": pick(["BID", "TID", "QD", "Q6H"]),
                    "start_datetime": start_base.strftime("%Y-%m-%d %H:%M:%S"),
                    "end_datetime": end_dt.strftime("%Y-%m-%d %H:%M:%S"),
                })
                prescription_id += 1

        # fill remainder
        while len(prescription_rows) < prescription_count:
            row = hosp_df.sample(n=1).iloc[0]
            safe = safe_drugs_for_patient(row.patient_amka)
            if not safe:
                continue
            start_base = datetime.fromisoformat(row.admission_ts) + timedelta(hours=random.randint(1, 48))
            end_limit = datetime.fromisoformat(row.discharge_ts)
            end_dt = min(end_limit, start_base + timedelta(days=random.randint(1, 7)))
            if end_dt <= start_base:
                continue
            doctor_choices = hosp_doc_df.loc[hosp_doc_df["hosp_id"] == row.hosp_id, "doctor_amka"].tolist()
            prescription_rows.append({
                "prescription_id": prescription_id,
                "hosp_id": int(row.hosp_id),
                "patient_amka": row.patient_amka,
                "doctor_amka": pick(doctor_choices),
                "drug_id": pick(safe),
                "dosage": pick(["1 tablet", "500 mg", "1 vial", "1 capsule", "40 mg"]),
                "frequency": pick(["BID", "TID", "QD", "Q6H", "Q8H"]),
                "start_datetime": start_base.strftime("%Y-%m-%d %H:%M:%S"),
                "end_datetime": end_dt.strftime("%Y-%m-%d %H:%M:%S"),
            })
            prescription_id += 1
        presc_df = pd.DataFrame(prescription_rows).drop_duplicates(subset=["doctor_amka", "patient_amka", "drug_id", "start_datetime"])
        seen_prescriptions = set(
            presc_df[["doctor_amka", "patient_amka", "drug_id", "start_datetime"]]
            .astype(str)
            .itertuples(index=False, name=None)
        )
        prescription_rows = presc_df.to_dict("records")
        attempts = 0
        while len(prescription_rows) < prescription_count and attempts < prescription_count * 20:
            attempts += 1
            row = hosp_df.sample(n=1, random_state=SEED + attempts).iloc[0]
            safe = safe_drugs_for_patient(row.patient_amka)
            if not safe:
                continue
            start_base = datetime.fromisoformat(row.admission_ts) + timedelta(hours=random.randint(1, 48), minutes=attempts % 60)
            end_limit = datetime.fromisoformat(row.discharge_ts)
            end_dt = min(end_limit, start_base + timedelta(days=random.randint(1, 7)))
            if end_dt <= start_base:
                continue
            doctor_choices = hosp_doc_df.loc[hosp_doc_df["hosp_id"] == row.hosp_id, "doctor_amka"].tolist()
            drug_id = pick(safe)
            key = (str(pick(doctor_choices)), str(row.patient_amka), str(drug_id), start_base.strftime("%Y-%m-%d %H:%M:%S"))
            if key in seen_prescriptions:
                continue
            seen_prescriptions.add(key)
            prescription_rows.append({
                "prescription_id": prescription_id,
                "hosp_id": int(row.hosp_id),
                "patient_amka": row.patient_amka,
                "doctor_amka": key[0],
                "drug_id": drug_id,
                "dosage": pick(["1 tablet", "500 mg", "1 vial", "1 capsule", "40 mg"]),
                "frequency": pick(["BID", "TID", "QD", "Q6H", "Q8H"]),
                "start_datetime": key[3],
                "end_datetime": end_dt.strftime("%Y-%m-%d %H:%M:%S"),
            })
            prescription_id += 1
        presc_df = pd.DataFrame(prescription_rows).sort_values("prescription_id").reset_index(drop=True)
        minimum_expected_prescriptions = min(360, prescription_count)
        if len(presc_df) < minimum_expected_prescriptions:
            raise ValueError(f"Could not generate {minimum_expected_prescriptions} unique, allergy-safe prescriptions.")

    # Minimal images
    for img_id, dep in enumerate(dept_df.itertuples(index=False), start=1):
        image_rows.append({
            "image_id": img_id,
            "image_url": f"https://example.com/images/department_{dep.department_id}.jpg",
            "alt_text": f"{dep.department_name} department photo placeholder",
        })
        entity_image_rows.append({
            "entity_name": "department",
            "entity_pk": str(dep.department_id),
            "image_id": img_id,
            "entity_description": f"Representative placeholder image for {dep.department_name}",
        })
    image_df = pd.DataFrame(image_rows)
    entity_image_df = pd.DataFrame(entity_image_rows)

    outputs = {
        "hospitalization": hosp_df,
        "hospitalization_doctor": hosp_doc_df,
        "lab_test": lab_df,
        "procedure_event": proc_event_df,
        "procedure_participant": proc_part_df,
        "patient_allergy": allergy_df,
        "prescription": presc_df,
        "hospitalization_evaluation": eval_df,
        "image_asset": image_df,
        "entity_image": entity_image_df,
    }
    for name, df in outputs.items():
        df.to_csv(gen_dir / f"{name}.csv", index=False)

    return outputs


def validate_generated_bundle(ref, org, patients, generated):
    """Fail fast if generated clinical rows drift from the reference CSVs."""

    def require_subset(label, values, allowed):
        missing = sorted(set(pd.Series(values).dropna().astype(str)) - set(pd.Series(allowed).dropna().astype(str)))
        if missing:
            preview = ", ".join(missing[:10])
            raise ValueError(f"{label} contains values missing from its reference table: {preview}")

    hospitalization = generated["hospitalization"]
    procedure_event = generated["procedure_event"]
    prescription = generated["prescription"]
    patient_allergy = generated["patient_allergy"]

    require_subset("hospitalization.ken_code", hospitalization["ken_code"], ref["ken"]["ken_code"])
    require_subset("icd10_ken_map.ken_code", ref["icd10_ken"]["ken_code"], ref["ken"]["ken_code"])
    require_subset("hospitalization.admission_icd10_code", hospitalization["admission_icd10_code"], ref["icd10"]["icd10_code"])
    require_subset("hospitalization.discharge_icd10_code", hospitalization["discharge_icd10_code"], ref["icd10"]["icd10_code"])
    require_subset("lab_test.test_code", generated["lab_test"]["test_code"], ref["lab_test_catalog"]["test_code"])

    require_subset("procedure_event.procedure_code", procedure_event["procedure_code"], ref["procedure_catalog"]["procedure_code"])
    proc_place = procedure_event.merge(
        ref["procedure_catalog"][["procedure_code", "required_place_type"]],
        on="procedure_code",
        how="left",
    ).merge(
        org["operating_place"][["place_id", "place_type"]],
        on="place_id",
        how="left",
    )
    bad_place = proc_place[proc_place["required_place_type"] != proc_place["place_type"]]
    if not bad_place.empty:
        raise ValueError("procedure_event.place_id does not match procedure_catalog.required_place_type.")

    require_subset("prescription.drug_id", prescription["drug_id"], ref["drug"]["drug_id"])
    require_subset("patient_allergy.substance_id", patient_allergy["substance_id"], ref["active_substance"]["substance_id"])
    require_subset("drug_active_substance.drug_id", ref["drug_active_substance"]["drug_id"], ref["drug"]["drug_id"])
    require_subset("drug_active_substance.substance_id", ref["drug_active_substance"]["substance_id"], ref["active_substance"]["substance_id"])
    require_subset("prescription.patient_amka", prescription["patient_amka"], patients["patient"]["patient_amka"])
    require_subset("patient_allergy.patient_amka", patient_allergy["patient_amka"], patients["patient"]["patient_amka"])


def write_load_sql(bundle_dir: Path):
    """Write a robust MySQL/MariaDB loader with explicit column lists.

    The order matches foreign-key dependencies and the CSV columns generated by this script.
    Nullable FK/date fields are loaded through user variables and converted with NULLIF.
    """
    sql_dir = bundle_dir / "sql"
    sql_dir.mkdir(exist_ok=True)

    def block(path, table, cols, set_lines=None):
        lines = [
            f"LOAD DATA LOCAL INFILE '{path}'",
            f"INTO TABLE {table}",
            "CHARACTER SET utf8mb4",
            "FIELDS TERMINATED BY ',' OPTIONALLY ENCLOSED BY '\"'",
            "LINES TERMINATED BY '\\n'",
            "IGNORE 1 LINES",
            "(" + ", ".join(cols) + ")"
        ]
        if set_lines:
            lines.append("SET")
            lines.extend(["    " + s + ("," if i < len(set_lines)-1 else "") for i, s in enumerate(set_lines)])
        lines[-1] += ";"
        return lines + [""]

    load_lines = [
        "-- Robust generated loader. Run from the bundle/project root.",
        "-- Example: mysql --local-infile=1 -u root -p < sql/load.sql",
        "USE yg_eupolis_hospital;",
        "SET FOREIGN_KEY_CHECKS = 1;",
        "",
    ]

    load_lines += block('data/reference/icd10_diagnosis.csv', 'icd10_diagnosis', ['icd10_code','icd10_description'])
    load_lines += block('data/reference/ken.csv', 'ken', ['ken_code','ken_description','basic_cost','mean_duration_days','extra_daily_cost'])
    load_lines += block('data/reference/procedure_catalog.csv', 'procedure_catalog', ['procedure_code','procedure_name','procedure_category','required_place_type'])
    load_lines += block('data/reference/lab_test_catalog.csv', 'lab_test_catalog', ['test_code','test_name','test_type'])
    load_lines += block('data/reference/drug.csv', 'drug', ['drug_id','drug_name'])
    load_lines += block('data/reference/active_substance.csv', 'active_substance', ['substance_id','substance_name'])
    load_lines += block('data/reference/drug_active_substance.csv', 'drug_active_substance', ['drug_id','substance_id'])

    # personnel must load before doctor; doctor must load before department because department.manager_doctor_amka references doctor.
    load_lines += block('data/generated/personnel.csv', 'personnel', ['amka','first_name','last_name','age','email','phone_number','hiring_date','personnel_type'])
    load_lines += block('data/generated/doctor.csv', 'doctor', ['@amka','@license_number','@specialization','@doctor_rank','@supervisor_amka'], [
        "amka = @amka",
        "license_number = @license_number",
        "specialization = @specialization",
        "doctor_rank = @doctor_rank",
        "supervisor_amka = NULLIF(@supervisor_amka, '')",
    ])
    load_lines += block('data/generated/department.csv', 'department', ['@department_id','@department_name','@description','@bed_capacity','@floor_building','@manager_doctor_amka'], [
        "department_id = @department_id",
        "department_name = @department_name",
        "description = NULLIF(@description, '')",
        "bed_capacity = @bed_capacity",
        "floor_building = @floor_building",
        "manager_doctor_amka = NULLIF(@manager_doctor_amka, '')",
    ])
    load_lines += block('data/generated/doctor_department.csv', 'doctor_department', ['doctor_amka','department_id'])
    load_lines += block('data/generated/nurse.csv', 'nurse', ['amka','nurse_rank','department_id'])
    load_lines += block('data/generated/administrative_staff.csv', 'administrative_staff', ['amka','admin_role','office_work','department_id'])
    load_lines += block('data/generated/bed.csv', 'bed', ['bed_id','department_id','bed_type','bed_status'])
    load_lines += block('data/generated/operating_place.csv', 'operating_place', ['place_id','place_name','place_type','place_status'])
    load_lines += block('data/generated/patient.csv', 'patient', ['@patient_amka','@first_name','@last_name','@father_name','@age','@gender','@weight_kg','@height_cm','@address_line','@phone_number','@email','@profession','@nationality','@insurance_provider'], [
        "patient_amka = @patient_amka",
        "first_name = @first_name",
        "last_name = @last_name",
        "father_name = @father_name",
        "age = @age",
        "gender = @gender",
        "weight_kg = NULLIF(@weight_kg, '')",
        "height_cm = NULLIF(@height_cm, '')",
        "address_line = @address_line",
        "phone_number = @phone_number",
        "email = NULLIF(@email, '')",
        "profession = NULLIF(@profession, '')",
        "nationality = NULLIF(@nationality, '')",
        "insurance_provider = @insurance_provider",
    ])
    load_lines += block('data/generated/emergency_contact.csv', 'emergency_contact', ['@patient_amka','@first_name','@last_name','@phone_number','@email'], [
        "patient_amka = @patient_amka",
        "first_name = @first_name",
        "last_name = @last_name",
        "phone_number = @phone_number",
        "email = NULLIF(@email, '')",
    ])
    load_lines += block('data/generated/department_shift.csv', 'department_shift', ['shift_id','department_id','shift_date','shift_type','start_time','end_time','shift_status'])
    load_lines += block('data/generated/shift_assignment.csv', 'shift_assignment', ['@shift_id','@personnel_amka','@assigned_role'], [
        "shift_id = @shift_id",
        "personnel_amka = @personnel_amka",
        "assigned_role = NULLIF(@assigned_role, '')",
    ])
    load_lines += [
        "-- Mark shifts as valid only after all staff assignments have loaded.",
        "-- This activates the vol2 shift-composition and resident-supervisor checks.",
        "UPDATE department_shift",
        "SET shift_status = 'VALID'",
        "WHERE shift_status = 'PROCESSING';",
        "",
    ]
    load_lines += block('data/generated/emergency_visit.csv', 'emergency_visit', ['@visit_id','@patient_amka','@triage_nurse_amka','@arrival_ts','@symptoms','@emergency_level','@service_start_ts','@disposition','@referred_department_id','@discharge_instructions','@status'], [
        "visit_id = @visit_id",
        "patient_amka = @patient_amka",
        "triage_nurse_amka = @triage_nurse_amka",
        "arrival_ts = @arrival_ts",
        "symptoms = @symptoms",
        "emergency_level = @emergency_level",
        "service_start_ts = NULLIF(@service_start_ts, '')",
        "disposition = @disposition",
        "referred_department_id = NULLIF(@referred_department_id, '')",
        "discharge_instructions = NULLIF(@discharge_instructions, '')",
        "status = NULLIF(@status, '')",
    ])
    load_lines += block('data/generated/hospitalization.csv', 'hospitalization', ['@hosp_id','@patient_amka','@department_id','@bed_id','@ken_code','@admission_ts','@discharge_ts','@admission_icd10_code','@discharge_icd10_code','@total_cost'], [
        "hosp_id = @hosp_id",
        "patient_amka = @patient_amka",
        "department_id = @department_id",
        "bed_id = @bed_id",
        "ken_code = @ken_code",
        "admission_ts = @admission_ts",
        "discharge_ts = NULLIF(@discharge_ts, '')",
        "admission_icd10_code = @admission_icd10_code",
        "discharge_icd10_code = NULLIF(@discharge_icd10_code, '')",
        "total_cost = @total_cost",
    ])
    load_lines += block('data/generated/hospitalization_doctor.csv', 'hospitalization_doctor', ['hosp_id','doctor_amka'])
    load_lines += block('data/generated/lab_test.csv', 'lab_test', ['@test_id','@hosp_id','@test_code','@ordered_by_doctor_amka','@test_datetime','@result_text'], [
        "test_id = @test_id",
        "hosp_id = @hosp_id",
        "test_code = @test_code",
        "ordered_by_doctor_amka = @ordered_by_doctor_amka",
        "test_datetime = @test_datetime",
        "result_text = NULLIF(@result_text, '')",
    ])
    load_lines += block('data/generated/procedure_event.csv', 'procedure_event', ['procedure_event_id','hosp_id','procedure_code','place_id','chief_surgeon_amka','start_ts','end_ts','actual_duration_min'])
    load_lines += block('data/generated/procedure_participant.csv', 'procedure_participant', ['procedure_event_id','personnel_amka'])
    load_lines += block('data/generated/patient_allergy.csv', 'patient_allergy', ['patient_amka','substance_id'])
    load_lines += block('data/generated/prescription.csv', 'prescription', ['@prescription_id','@hosp_id','@patient_amka','@doctor_amka','@drug_id','@dosage','@frequency','@start_datetime','@end_datetime'], [
        "prescription_id = @prescription_id",
        "hosp_id = @hosp_id",
        "patient_amka = @patient_amka",
        "doctor_amka = @doctor_amka",
        "drug_id = @drug_id",
        "dosage = @dosage",
        "frequency = @frequency",
        "start_datetime = @start_datetime",
        "end_datetime = NULLIF(@end_datetime, '')",
    ])
    load_lines += block('data/generated/hospitalization_evaluation.csv', 'hospitalization_evaluation', ['@hosp_id','@evaluation_date','@medical_care_score','@nursing_care_score','@cleanliness_score','@food_score','@overall_experience_score','@comments'], [
        "hosp_id = @hosp_id",
        "evaluation_date = @evaluation_date",
        "medical_care_score = @medical_care_score",
        "nursing_care_score = @nursing_care_score",
        "cleanliness_score = @cleanliness_score",
        "food_score = @food_score",
        "overall_experience_score = @overall_experience_score",
        "comments = NULLIF(@comments, '')",
    ])
    load_lines += block('data/generated/image_asset.csv', 'image_asset', ['image_id','image_url','alt_text'])
    load_lines += block('data/generated/entity_image.csv', 'entity_image', ['@entity_name','@entity_pk','@image_id','@entity_description'], [
        "entity_name = @entity_name",
        "entity_pk = @entity_pk",
        "image_id = @image_id",
        "entity_description = NULLIF(@entity_description, '')",
    ])

    (sql_dir / "load.sql").write_text("\n".join(load_lines), encoding="utf-8")

def write_guides(bundle_dir: Path, ref, generated):
    # table map
    rows = []
    for csv in sorted((bundle_dir / "data/reference").glob("*.csv")):
        table = csv.stem
        rows.append({"table_name": table, "csv_path": f"data/reference/{csv.name}", "kind": "reference"})
    for csv in sorted((bundle_dir / "data/generated").glob("*.csv")):
        table = csv.stem
        rows.append({"table_name": table, "csv_path": f"data/generated/{csv.name}", "kind": "synthetic"})
    pd.DataFrame(rows).to_csv(bundle_dir / "TABLE_TO_CSV_MAP.csv", index=False)

    coverage = [
        {"query_id":"Q1","seed_rule":"Hospitalizations cover 2025-2026, multiple KEN codes, multiple insurance providers, and many stays exceed MDN."},
        {"query_id":"Q2","seed_rule":"Procedure events are skewed toward a set of senior surgeons; all doctors belong to at least one department and one week of shifts exists."},
        {"query_id":"Q3","seed_rule":"15 patients have 4-5 hospitalizations in the same department."},
        {"query_id":"Q4","seed_rule":"~74% of discharged hospitalizations have an evaluation; each hospitalization has at least one linked doctor."},
        {"query_id":"Q5","seed_rule":"Young doctors (<35) exist and some of them are assigned chief-surgeon procedure events."},
        {"query_id":"Q6","seed_rule":"Several patients have long hospitalization histories with ICD-10, total cost and evaluation coverage."},
        {"query_id":"Q7","seed_rule":"Drug-substance links and patient allergies are populated, with popular substances appearing often."},
        {"query_id":"Q8","seed_rule":"A full one-week staffing schedule exists for all 15 departments and all 3 daily shifts."},
        {"query_id":"Q9","seed_rule":"10 patient-pairs have equal total hospitalization days in 2026 and totals > 15 days."},
        {"query_id":"Q10","seed_rule":"Three active-substance pairs are planted repeatedly in prescriptions within the same hospitalization."},
        {"query_id":"Q11","seed_rule":"Chief-surgeon counts are deliberately uneven across doctors."},
        {"query_id":"Q12","seed_rule":"Shifts include subtype detail (doctor specialization, nurse rank, admin role) for a specific week."},
        {"query_id":"Q13","seed_rule":"Doctor supervision is generated top-down by rank without cycles."},
        {"query_id":"Q14","seed_rule":"Five ICD-10 3-character categories are seeded with equal counts in 2025 and 2026, each >= 6 admissions/year."},
        {"query_id":"Q15","seed_rule":"Emergency visits cover levels 1-5 with waiting times, hospitalization disposition and referral departments."},
    ]
    pd.DataFrame(coverage).to_csv(bundle_dir / "QUERY_COVERAGE.csv", index=False)

    notes = f"""# Dataset guide

This bundle contains:

- cleaned **reference CSVs** from the uploaded official files;
- a deterministic **query-driven synthetic dataset**;
    - a Python generator script: `scripts/generate_data.py`;
- a schema installer copied from the project: `sql/install.sql`;
- a convenience loader: `sql/load.sql`;
- a validation script copied from the project: `sql/validation.sql`;
- a table-to-CSV manifest: `TABLE_TO_CSV_MAP.csv`.

## Important schema assumptions

This dataset targets the **Ygeiopolis vol2 schema**, with the following minimal clarifications so the data load stays coherent:

1. `nurse` uses `nurse_rank` (not `degree`).
2. `department_shift` is generated as `PROCESSING` and the loader marks it `VALID` after staff assignments load, so the vol2 shift procedures validate coverage.
3. `emergency_visit` includes `referred_department_id` and the vol2 `status` field.
4. `hospitalization_doctor` stores the doctor link only, matching vol2.
5. `procedure_participant` stores the participant link only, matching vol2.
6. `ken.csv` uses only the improved/official KEN export; the generator fails clearly if official KEN data cannot be found or parsed.
7. If the official EMA Article 57 workbook is not supplied, drug/prescription/allergy CSVs remain empty by default so the final load contains no unofficial medication data.
8. Procedure codes and names come from the official procedure catalog; vol2 keeps category and required place type in the loaded table.

## Name mapping logic

- CSV file name == table name whenever possible.
- Official source -> cleaned output:
  - `4.2 Κωδικοί ICD-10...xls` -> `data/reference/icd10_diagnosis.csv`
  - improved/official KEN export or official KEN `.doc` -> `data/reference/ken.csv`
  - official ICD10-KEN workbook, filtered against `ken.csv` -> `data/reference/icd10_ken_map.csv`
  - `ΕΛΛΗΝΙΚΗ ΟΝΟΜΑΤΟΛΟΓΙΑ ... ΙΑΤΡΙΚΩΝ ΠΡΑΞΕΩΝ...xls` -> `data/reference/procedure_catalog.csv`
  - same procedure workbook -> derived `data/reference/lab_test_catalog.csv`

## How to rerun

```bash
python scripts/generate_data.py --source-dir . --output-dir .
```

To use official EMA Article 57 data once you have it:

```bash
python scripts/generate_data.py --source-dir . --output-dir . --ema-xlsx /path/to/article-57-product-data_en.xlsx
```

"""
    (bundle_dir / "DATASET_GUIDE.md").write_text(notes, encoding="utf-8")


def main():
    parser = argparse.ArgumentParser(description="Clean reference files and generate a query-driven synthetic dataset.")
    script_path = Path(__file__).resolve()
    default_root = script_path.parents[1] if script_path.parent.name == "scripts" else script_path.parent
    parser.add_argument("--source-dir", default=str(default_root), help="Directory containing the uploaded source files.")
    parser.add_argument("--output-dir", default=str(default_root / "hospital_dataset_bundle"), help="Output bundle root.")
    parser.add_argument("--ema-xlsx", default=None, help="Optional EMA Article 57 workbook for official drug reference data.")
    parser.add_argument("--patient-count", type=int, default=5000, help="Number of synthetic patients to generate.")
    parser.add_argument("--emergency-count", type=int, default=18000, help="Number of synthetic emergency visits to generate.")
    parser.add_argument("--hospitalization-count", type=int, default=12000, help="Target number of hospitalizations to generate.")
    parser.add_argument("--lab-test-count", type=int, default=12000, help="Target number of lab tests to generate.")
    parser.add_argument("--procedure-count", type=int, default=6000, help="Target number of procedure events to attempt.")
    parser.add_argument("--prescription-count", type=int, default=12000, help="Target number of prescriptions to generate.")
    parser.add_argument("--shift-start-date", default=DATASET_START_DATE.isoformat(), help="First generated shift date, YYYY-MM-DD.")
    parser.add_argument("--shift-days", type=int, default=DEFAULT_SHIFT_DAYS, help="Operational date window in days for shifts, emergency visits, and hospitalizations.")
    parser.add_argument("--shift-sample-days-per-month", type=int, default=DEFAULT_SHIFT_SAMPLE_DAYS_PER_MONTH, help="Number of fully covered shift days generated per month; use 0 for every day.")
    args = parser.parse_args()

    source_dir = Path(args.source_dir).expanduser().resolve()
    bundle_dir = Path(args.output_dir).expanduser().resolve()
    shift_start_date = date.fromisoformat(args.shift_start_date)
    ref_dir = bundle_dir / "data" / "reference"
    gen_dir = bundle_dir / "data" / "generated"
    (bundle_dir / "data" / "reference").mkdir(parents=True, exist_ok=True)
    (bundle_dir / "data" / "generated").mkdir(parents=True, exist_ok=True)
    (bundle_dir / "scripts").mkdir(parents=True, exist_ok=True)

    ref = load_reference_data(source_dir, ref_dir, ema_xlsx=Path(args.ema_xlsx) if args.ema_xlsx else None)
    org = build_people_and_org(gen_dir)
    patients = build_patients(gen_dir, count=args.patient_count)
    shifts = build_shifts(
        gen_dir,
        org,
        start_date=shift_start_date,
        days=args.shift_days,
        sample_days_per_month=args.shift_sample_days_per_month,
    )
    emergencies = build_emergency_visits(gen_dir, patients, org, count=args.emergency_count, start_date=shift_start_date, days=args.shift_days)
    clinical = build_clinical_events(
        gen_dir,
        ref,
        org,
        patients,
        hospitalization_count=args.hospitalization_count,
        lab_test_count=args.lab_test_count,
        procedure_event_count=args.procedure_count,
        prescription_count=args.prescription_count,
        admission_start_date=shift_start_date,
        admission_days=args.shift_days,
    )
    validate_generated_bundle(ref, org, patients, clinical)
    write_load_sql(bundle_dir)
    project_sql_dir = script_path.parents[1] / "sql"
    for sql_name in ["install.sql", "validation.sql"]:
        source_sql = project_sql_dir / sql_name
        target_sql = bundle_dir / "sql" / sql_name
        if source_sql.exists() and source_sql.resolve() != target_sql.resolve():
            shutil.copy2(source_sql, target_sql)
    write_guides(bundle_dir, ref, {**org, **patients, **shifts, **emergencies, **clinical})
    target_script = bundle_dir / "scripts" / "generate_data.py"
    if script_path.resolve() != target_script.resolve():
        shutil.copy2(script_path, target_script)

    summary = {
        "reference_rows": {k: int(len(v)) for k, v in ref.items() if isinstance(v, pd.DataFrame)},
        "synthetic_rows": {
            "department": int(len(org["department"])),
            "personnel": int(len(org["personnel"])),
            "doctor": int(len(org["doctor"])),
            "nurse": int(len(org["nurse"])),
            "administrative_staff": int(len(org["administrative_staff"])),
            "bed": int(len(org["bed"])),
            "patient": int(len(patients["patient"])),
            "emergency_contact": int(len(patients["emergency_contact"])),
            "department_shift": int(len(shifts["department_shift"])),
            "shift_assignment": int(len(shifts["shift_assignment"])),
            "emergency_visit": int(len(emergencies["emergency_visit"])),
            "hospitalization": int(len(clinical["hospitalization"])),
            "hospitalization_doctor": int(len(clinical["hospitalization_doctor"])),
            "lab_test": int(len(clinical["lab_test"])),
            "procedure_event": int(len(clinical["procedure_event"])),
            "procedure_participant": int(len(clinical["procedure_participant"])),
            "patient_allergy": int(len(clinical["patient_allergy"])),
            "prescription": int(len(clinical["prescription"])),
            "hospitalization_evaluation": int(len(clinical["hospitalization_evaluation"])),
        },
        "ema_mode": ref["meta"]["ema_mode"],
    }
    (bundle_dir / "dataset_summary.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
